"""
Convert all markdown files in outputs/ to properly formatted Word documents.

Usage:
    python scripts/md_to_docx.py
    python scripts/md_to_docx.py --dir outputs --zip

Handles:
  # H1  ## H2  ### H3
  **bold**  *italic*  ***bold+italic***  `code`
  - bullet lists
  1. numbered lists
  > blockquotes
  --- horizontal rule
"""
from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Encoding cleanup ──────────────────────────────────────────────────────────
# Mojibake family A: UTF-8 smart-punctuation bytes decoded as Latin-1
#   e.g. RIGHT SINGLE QUOTE U+2019 has UTF-8 bytes \xe2\x80\x99
#        Decoded as Latin-1 those become U+00E2, U+0080, U+0099
#
# Mojibake family B: same bytes decoded as Windows-1252 (CP1252)
#   \x80 -> U+20AC (euro), \x99 -> U+2122 (TM) etc.
#
# Longer patterns must come first so they shadow shorter overlapping ones.
_MOJIBAKE = [
    # ── Latin-1 decoded (family A) ───────────────────────────────────────────
    # right single quote U+2019: bytes \xe2\x80\x99
    ("â", "’"),
    # left  single quote U+2018: bytes \xe2\x80\x98
    ("â", "‘"),
    # right double quote U+201D: bytes \xe2\x80\x9d
    ("â", "”"),
    # left  double quote U+201C: bytes \xe2\x80\x9c
    ("â", "“"),
    # em-dash U+2014: bytes \xe2\x80\x94
    ("â", "—"),
    # en-dash U+2013: bytes \xe2\x80\x93
    ("â", "–"),
    # ellipsis U+2026: bytes \xe2\x80\xa6
    ("â¦", "…"),
    # bullet U+2022: bytes \xe2\x80\xa2 (less common but possible)
    ("â¢", "•"),
    # ── CP1252 decoded (family B) ─────────────────────────────────────────────
    # right single quote: â€™  (â + euro + TM)
    ("â€™", "’"),
    # left  double quote: â€œ  (â + euro + LATIN SMALL LETTER Z WITH CARON)
    ("â€œ", "“"),
    # right double quote: â€   (â + euro + right double quotation)
    ("â€", "”"),
    # em-dash: â€"  (â + euro + en-dash)
    ("â€”", "—"),
    # en-dash: â€"  (â + euro + em-dash)
    ("â€“", "–"),
    # ellipsis: â€¦  (â + euro + ¦)
    ("â€¦", "…"),
    # left  single quote: â€˜  (â + euro + LATIN SMALL LETTER Z WITH ACUTE)
    ("â€˜", "‘"),
]


def _clean(text: str) -> str:
    """Fix mojibake. Keeps Unicode dashes and quotes as proper Unicode."""
    text = text.lstrip("﻿")  # strip BOM
    for bad, good in _MOJIBAKE:
        text = text.replace(bad, good)
    return text


# ── Inline markdown -> runs ───────────────────────────────────────────────────
_INLINE = re.compile(
    r"\*\*\*(.+?)\*\*\*"   # ***bold+italic***
    r"|\*\*(.+?)\*\*"       # **bold**
    r"|\*(.+?)\*"           # *italic*
    r"|`(.+?)`"             # `code`
)


def _add_inline(para, text: str):
    """Add text with inline bold/italic/code formatting to a paragraph."""
    text = _clean(text)
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            r = para.add_run(m.group(1)); r.bold = True; r.italic = True
        elif m.group(2):
            r = para.add_run(m.group(2)); r.bold = True
        elif m.group(3):
            r = para.add_run(m.group(3)); r.italic = True
        elif m.group(4):
            r = para.add_run(m.group(4)); r.font.name = "Courier New"; r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _add_hr(doc):
    """Add a thin horizontal line (mimics --- in markdown)."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Main converter ─────────────────────────────────────────────────────────────
def markdown_to_docx(title: str, markdown: str, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # Document title
    tp = doc.add_heading(_clean(title), level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            h = re.sub(r"\*+(.+?)\*+", r"\1", _clean(stripped[4:]))
            doc.add_heading(h, level=3)
        elif stripped.startswith("## "):
            h = re.sub(r"\*+(.+?)\*+", r"\1", _clean(stripped[3:]))
            doc.add_heading(h, level=2)
        elif stripped.startswith("# "):
            h = re.sub(r"\*+(.+?)\*+", r"\1", _clean(stripped[2:]))
            doc.add_heading(h, level=1)
        # Horizontal rule — only when entire stripped line is dashes or asterisks
        elif re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            _add_hr(doc)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, stripped[2:])
        elif re.match(r"^\d+[\.\)]\s", stripped):
            text = re.sub(r"^\d+[\.\)]\s+", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, text)
        elif stripped.startswith("> "):
            para = doc.add_paragraph(style="Quote")
            _add_inline(para, stripped[2:])
        else:
            para = doc.add_paragraph()
            _add_inline(para, stripped)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ── CLI ────────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--dir", default="outputs", help="Directory containing .md files")
    parser.add_argument("--zip", action="store_true", help="Also create a zip of all .docx files")
    args = parser.parse_args()

    src = Path(args.dir)
    md_files = sorted(src.glob("*.md"))

    if not md_files:
        print(f"No .md files found in {src}/")
        return

    docx_files = []
    skipped = []
    for md_path in md_files:
        markdown = md_path.read_text(encoding="utf-8-sig").strip()
        words = len(markdown.split())

        if words < 50:
            skipped.append(md_path.name)
            continue
        if md_path.stem.startswith("draft_"):
            skipped.append(md_path.name)
            continue

        title = md_path.stem.replace("_", " ").strip()
        title = re.sub(r"^\d+\s+", "", title)

        out_path = md_path.with_suffix(".docx")
        markdown_to_docx(title, markdown, out_path)
        docx_files.append(out_path)
        print(f"  OK  {out_path.name}  ({words} words)")

    if skipped:
        print(f"\nSkipped {len(skipped)} file(s): {', '.join(skipped)}")

    print(f"\nConverted {len(docx_files)} files.")

    if args.zip and docx_files:
        from datetime import datetime
        zip_path = src / f"articles_docx_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in docx_files:
                zf.write(f, f.name)
        print(f"Zip saved: {zip_path}  ({len(docx_files)} docs)")


if __name__ == "__main__":
    main()
