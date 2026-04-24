"""
Sequential runner for the article pipeline endpoints.

Flow:
1) Login (unless token provided)
2) POST /api/v1/articles/pipeline
3) Poll GET /api/v1/articles/pipeline/{pipeline_id}
4) If failed, PATCH /api/v1/articles/pipeline/{pipeline_id}/resume (optional)
5) Poll again until terminal state

Built with:
- Request retry for transient failures
- Per-request timeout (connect/read)
- Global timeout
- Stall/freeze timeout (no state change for too long)
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests


TERMINAL_STATES = {"completed", "completed_with_warnings", "failed"}


def _now() -> float:
    return time.monotonic()


def _fmt_keywords(value: str) -> List[str]:
    if not value:
        return []
    parts = [p.strip() for p in value.split(",")]
    return [p for p in parts if p]


def _request_with_retries(
    session: requests.Session,
    method: str,
    url: str,
    *,
    headers: Optional[Dict[str, str]] = None,
    payload: Optional[Dict[str, Any]] = None,
    timeout: Tuple[float, float] = (8.0, 60.0),
    max_retries: int = 3,
    backoff_seconds: float = 1.0,
) -> requests.Response:
    last_error: Optional[Exception] = None
    for attempt in range(max_retries + 1):
        try:
            resp = session.request(
                method=method,
                url=url,
                headers=headers,
                json=payload,
                timeout=timeout,
            )
        except (requests.Timeout, requests.ConnectionError) as exc:
            last_error = exc
            if attempt >= max_retries:
                break
            time.sleep(backoff_seconds * (2 ** attempt))
            continue

        if resp.status_code >= 500 and attempt < max_retries:
            time.sleep(backoff_seconds * (2 ** attempt))
            continue
        return resp

    if last_error:
        raise RuntimeError(f"Request failed after retries: {method} {url} :: {last_error}") from last_error
    raise RuntimeError(f"Request failed after retries: {method} {url}")


def _ensure_ok(resp: requests.Response, expected_status: int) -> Dict[str, Any]:
    if resp.status_code != expected_status:
        body = resp.text[:1200]
        raise RuntimeError(
            f"Unexpected status {resp.status_code} (expected {expected_status}). Body: {body}"
        )
    try:
        return resp.json()
    except Exception as exc:
        raise RuntimeError(f"Response is not JSON: {resp.text[:600]}") from exc


def login(
    session: requests.Session,
    base_url: str,
    tenant_id: str,
    role: str,
    *,
    timeout: Tuple[float, float],
    max_retries: int,
) -> str:
    url = f"{base_url}/api/v1/auth/login"
    payload = {"tenant_id": tenant_id, "role": role}
    resp = _request_with_retries(
        session,
        "POST",
        url,
        payload=payload,
        timeout=timeout,
        max_retries=max_retries,
    )
    data = _ensure_ok(resp, 200)
    token = data.get("access_token")
    if not token:
        raise RuntimeError(f"Login response missing access_token: {data}")
    return str(token)


def start_pipeline(
    session: requests.Session,
    base_url: str,
    headers: Dict[str, str],
    body: Dict[str, Any],
    *,
    timeout: Tuple[float, float],
    max_retries: int,
) -> Dict[str, Any]:
    url = f"{base_url}/api/v1/articles/pipeline"
    resp = _request_with_retries(
        session,
        "POST",
        url,
        headers=headers,
        payload=body,
        timeout=timeout,
        max_retries=max_retries,
    )
    return _ensure_ok(resp, 202)


def get_pipeline_status(
    session: requests.Session,
    base_url: str,
    headers: Dict[str, str],
    pipeline_id: str,
    *,
    timeout: Tuple[float, float],
    max_retries: int,
) -> Dict[str, Any]:
    url = f"{base_url}/api/v1/articles/pipeline/{pipeline_id}"
    resp = _request_with_retries(
        session,
        "GET",
        url,
        headers=headers,
        timeout=timeout,
        max_retries=max_retries,
    )
    return _ensure_ok(resp, 200)


def resume_pipeline(
    session: requests.Session,
    base_url: str,
    headers: Dict[str, str],
    pipeline_id: str,
    *,
    timeout: Tuple[float, float],
    max_retries: int,
) -> Dict[str, Any]:
    url = f"{base_url}/api/v1/articles/pipeline/{pipeline_id}/resume"
    resp = _request_with_retries(
        session,
        "PATCH",
        url,
        headers=headers,
        payload={},
        timeout=timeout,
        max_retries=max_retries,
    )
    return _ensure_ok(resp, 200)


def poll_until_terminal(
    session: requests.Session,
    base_url: str,
    headers: Dict[str, str],
    pipeline_id: str,
    *,
    poll_interval: float,
    timeout: Tuple[float, float],
    max_retries: int,
    deadline_ts: float,
    stalled_timeout: float,
) -> Dict[str, Any]:
    last_marker: Optional[Tuple[str, str, str, str]] = None
    last_change_ts = _now()

    while True:
        if _now() >= deadline_ts:
            raise TimeoutError("Global timeout reached while waiting for pipeline completion.")

        data = get_pipeline_status(
            session,
            base_url,
            headers,
            pipeline_id,
            timeout=timeout,
            max_retries=max_retries,
        )

        status = str(data.get("pipeline_status") or "")
        step = str(data.get("current_step") or "")
        failed_step = str(data.get("failed_step") or "")
        error = str(data.get("error_detail") or "")
        marker = (status, step, failed_step, error)

        if marker != last_marker:
            last_marker = marker
            last_change_ts = _now()
            print(f"[poll] status={status} step={step} failed_step={failed_step}")
            if error:
                print(f"[poll] error_detail={error}")
        elif (_now() - last_change_ts) > stalled_timeout:
            raise TimeoutError(
                f"Pipeline appears stalled (no status/step change for {stalled_timeout:.0f}s)."
            )

        if status in TERMINAL_STATES:
            return data

        time.sleep(max(0.5, poll_interval))


def _slugify_filename(text: str) -> str:
    t = re.sub(r"[^\w\s-]", "", (text or "").strip(), flags=re.UNICODE)
    t = re.sub(r"[\s_-]+", "_", t).strip("_")
    return (t or "article").lower()


def _extract_markdown_title(markdown: str, fallback: str) -> str:
    for line in (markdown or "").splitlines():
        s = line.strip()
        if s.startswith("# "):
            return s[2:].strip() or fallback
    return fallback


def _inline_markdown_tokens(text: str) -> List[Tuple[str, str]]:
    """
    Returns list of (kind, value) where kind in:
    - "text"
    - "bold"
    - "italic"
    - "code"
    """
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    out: List[Tuple[str, str]] = []
    cursor = 0
    for m in pattern.finditer(text):
        if m.start() > cursor:
            out.append(("text", text[cursor : m.start()]))
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            out.append(("bold", chunk[2:-2]))
        elif chunk.startswith("*") and chunk.endswith("*"):
            out.append(("italic", chunk[1:-1]))
        elif chunk.startswith("`") and chunk.endswith("`"):
            out.append(("code", chunk[1:-1]))
        else:
            out.append(("text", chunk))
        cursor = m.end()
    if cursor < len(text):
        out.append(("text", text[cursor:]))
    return out


def _add_markdown_paragraph(doc: Any, text: str, style: Optional[str] = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for kind, value in _inline_markdown_tokens(text):
        run = p.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"


def _render_markdown_to_docx(doc: Any, markdown: str) -> None:
    for line in (markdown or "").splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph("")
            continue

        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2).strip(), level=level)
            continue

        if re.match(r"^[-*]\s+", s):
            _add_markdown_paragraph(doc, re.sub(r"^[-*]\s+", "", s), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", s):
            _add_markdown_paragraph(doc, re.sub(r"^\d+\.\s+", "", s), style="List Number")
            continue

        if s.startswith("> "):
            _add_markdown_paragraph(doc, s[2:], style="Intense Quote")
            continue

        if re.match(r"^---+$", s):
            doc.add_paragraph("")
            continue

        _add_markdown_paragraph(doc, s)


def _count_words_from_markdown(markdown: str) -> int:
    t = markdown or ""
    # Remove common markdown markers before counting
    t = re.sub(r"```[\s\S]*?```", " ", t)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    t = re.sub(r"(^|\n)\s{0,3}#{1,6}\s*", " ", t)
    t = re.sub(r"(^|\n)\s*[-*]\s+", " ", t)
    t = re.sub(r"(^|\n)\s*\d+\.\s+", " ", t)
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"\*([^*]+)\*", r"\1", t)
    words = re.findall(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?", t)
    return len(words)


def _save_final_to_docx(
    final_data: Dict[str, Any],
    *,
    requested_title: str,
    output_path: Optional[str] = None,
    strip_source_tags: bool = True,
) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception as exc:
        raise RuntimeError(
            "python-docx is required to save .docx output. Install with: pip install python-docx"
        ) from exc

    summary = final_data.get("result_summary") if isinstance(final_data.get("result_summary"), dict) else {}
    markdown = str(summary.get("article_markdown") or "")
    if strip_source_tags and markdown:
        markdown = re.sub(r"\s*\[S\d+\]", "", markdown)
        markdown = re.sub(r"\n{3,}", "\n\n", markdown).strip()
    final_word_count = _count_words_from_markdown(markdown)

    pipeline_id = str(final_data.get("pipeline_id") or "unknown")
    title = _extract_markdown_title(markdown, requested_title)

    if output_path:
        out = Path(output_path)
    else:
        slug = _slugify_filename(title)[:70]
        out = Path(f"{slug}_{pipeline_id[:8]}.docx")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(title, 0)
    subtitle = doc.add_paragraph()
    subtitle.add_run("Pipeline ID: ").bold = True
    subtitle.add_run(pipeline_id)
    subtitle.add_run("   ")
    subtitle.add_run("Status: ").bold = True
    subtitle.add_run(str(final_data.get("pipeline_status") or "unknown"))

    doc.add_heading("Run Summary", level=1)
    table = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    qc = summary.get("qc_metrics") if isinstance(summary.get("qc_metrics"), dict) else {}
    rows: List[Tuple[str, str]] = [
        ("Request ID", str(final_data.get("request_id") or "")),
        ("Started At", str(final_data.get("started_at") or "")),
        ("Completed At", str(final_data.get("completed_at") or "")),
        ("ZeroGPT Score", str(summary.get("zerogpt_score") or "")),
        ("ZeroGPT Pass", str(summary.get("zerogpt_pass") or "")),
        ("QC Pass", str(summary.get("qc_pass") or "")),
        ("Flesch Reading Ease", str(qc.get("flesch_reading_ease") or "")),
        ("Flesch-Kincaid Grade", str(qc.get("flesch_kincaid_grade") or "")),
        ("Word Count (Final Exported)", str(final_word_count)),
        ("Word Count (QC Snapshot)", str(qc.get("word_count") or "")),
        ("Total Tokens", str(summary.get("total_tokens") or "")),
        ("Warning", str(summary.get("warning") or "")),
    ]
    for field, value in rows:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = value

    doc.add_page_break()
    doc.add_heading("Article", level=1)
    _render_markdown_to_docx(doc, markdown)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out.resolve())


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Run article pipeline endpoints sequentially.")
    p.add_argument("--base-url", default=os.getenv("API_BASE_URL", "http://localhost:8000"))
    p.add_argument("--token", default=os.getenv("API_TOKEN", ""))
    p.add_argument("--tenant-id", default=os.getenv("TENANT_ID", "5ebc3c3a-5e0e-42a5-a350-76f1b792ac15"))
    p.add_argument("--role", default=os.getenv("ROLE", "tenant_admin"))

    p.add_argument("--pipeline-id", default="")
    p.add_argument("--kb-id", default="")
    p.add_argument("--title", default="")
    p.add_argument("--keywords", default="")
    p.add_argument("--length-target", type=int, default=1800)
    p.add_argument("--request-id", default="")

    p.add_argument("--skip-crawl", action=argparse.BooleanOptionalAction, default=True)
    p.add_argument("--url", default="")
    p.add_argument("--max-depth", type=int, default=0)
    p.add_argument("--max-pages", type=int, default=1)
    p.add_argument("--respect-robots", action=argparse.BooleanOptionalAction, default=False)
    p.add_argument(
        "--user-agent",
        default="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
    )

    p.add_argument("--draft-provider", default="openai")
    p.add_argument("--draft-model", default="")
    p.add_argument("--temperature", type=float, default=0.55)
    p.add_argument("--max-output-tokens", type=int, default=8192)
    p.add_argument("--top-k-sources", type=int, default=10)
    p.add_argument("--rag-grounding-ratio", type=float, default=0.95)
    p.add_argument("--enable-agentic-orchestration", action=argparse.BooleanOptionalAction, default=True)
    p.add_argument("--expanded-query-count", type=int, default=4)
    p.add_argument("--hybrid-top-k-per-query", type=int, default=30)
    p.add_argument("--predictability-top-n", type=int, default=14)
    p.add_argument("--max-predictability-rewrite-passes", type=int, default=1)
    p.add_argument("--zerogpt-fix-max-attempts", type=int, default=4)
    p.add_argument("--max-quality-retries", type=int, default=4)

    p.add_argument("--poll-interval", type=float, default=5.0)
    p.add_argument("--connect-timeout", type=float, default=8.0)
    p.add_argument("--read-timeout", type=float, default=60.0)
    p.add_argument("--max-retries", type=int, default=3)
    p.add_argument("--global-timeout", type=float, default=1800.0)
    p.add_argument("--stalled-timeout", type=float, default=420.0)
    p.add_argument("--resume-on-fail", action=argparse.BooleanOptionalAction, default=True)
    p.add_argument("--max-resume-attempts", type=int, default=1)
    p.add_argument("--save-docx", action=argparse.BooleanOptionalAction, default=False)
    p.add_argument("--docx-path", default="")
    p.add_argument("--strip-source-tags", action=argparse.BooleanOptionalAction, default=True)
    return p


def main() -> int:
    args = build_parser().parse_args()
    existing_pipeline_mode = bool((args.pipeline_id or "").strip())
    if not existing_pipeline_mode:
        if not args.kb_id or not args.title:
            print("ERROR: --kb-id and --title are required when starting a new pipeline.", file=sys.stderr)
            return 2
        if not args.skip_crawl and not args.url:
            print("ERROR: --url is required when --no-skip-crawl is used.", file=sys.stderr)
            return 2

    keywords = _fmt_keywords(args.keywords)
    timeout = (float(args.connect_timeout), float(args.read_timeout))
    deadline = _now() + float(args.global_timeout)
    body: Dict[str, Any] = {}
    if not existing_pipeline_mode:
        body = {
            "url": args.url or None,
            "skip_crawl": bool(args.skip_crawl),
            "kb_id": args.kb_id,
            "title": args.title,
            "keywords": keywords,
            "length_target": int(args.length_target),
            "max_depth": int(args.max_depth),
            "max_pages": int(args.max_pages),
            "respect_robots": bool(args.respect_robots),
            "user_agent": args.user_agent,
            "draft_provider": args.draft_provider,
            "draft_model": args.draft_model,
            "temperature": float(args.temperature),
            "max_output_tokens": int(args.max_output_tokens),
            "top_k_sources": int(args.top_k_sources),
            "rag_grounding_ratio": float(args.rag_grounding_ratio),
            "enable_agentic_orchestration": bool(args.enable_agentic_orchestration),
            "expanded_query_count": int(args.expanded_query_count),
            "hybrid_top_k_per_query": int(args.hybrid_top_k_per_query),
            "predictability_top_n": int(args.predictability_top_n),
            "max_predictability_rewrite_passes": int(args.max_predictability_rewrite_passes),
            "zerogpt_fix_max_attempts": int(args.zerogpt_fix_max_attempts),
            "max_quality_retries": int(args.max_quality_retries),
            "request_id": args.request_id or None,
        }

    with requests.Session() as session:
        token = (args.token or "").strip()
        if not token:
            print("[auth] logging in...")
            token = login(
                session,
                args.base_url.rstrip("/"),
                args.tenant_id,
                args.role,
                timeout=timeout,
                max_retries=int(args.max_retries),
            )
            print("[auth] token acquired.")
        else:
            print("[auth] using provided token.")

        headers = {"Authorization": f"Bearer {token}"}

        if existing_pipeline_mode:
            pipeline_id = str(args.pipeline_id).strip()
            print(f"[pipeline] watching existing pipeline_id={pipeline_id}")
        else:
            print("[pipeline] starting...")
            start_data = start_pipeline(
                session,
                args.base_url.rstrip("/"),
                headers,
                body,
                timeout=timeout,
                max_retries=int(args.max_retries),
            )
            pipeline_id = str(start_data.get("pipeline_id"))
            if not pipeline_id:
                print(f"ERROR: pipeline_id missing in response: {start_data}", file=sys.stderr)
                return 1
            print(f"[pipeline] started pipeline_id={pipeline_id}")

        final_data = poll_until_terminal(
            session,
            args.base_url.rstrip("/"),
            headers,
            pipeline_id,
            poll_interval=float(args.poll_interval),
            timeout=timeout,
            max_retries=int(args.max_retries),
            deadline_ts=deadline,
            stalled_timeout=float(args.stalled_timeout),
        )

        status = str(final_data.get("pipeline_status") or "")
        resume_count = 0
        while (
            status == "failed"
            and args.resume_on_fail
            and resume_count < int(args.max_resume_attempts)
            and _now() < deadline
        ):
            resume_count += 1
            print(f"[pipeline] failed -> resume attempt {resume_count}...")
            resume_resp = resume_pipeline(
                session,
                args.base_url.rstrip("/"),
                headers,
                pipeline_id,
                timeout=timeout,
                max_retries=int(args.max_retries),
            )
            print(f"[pipeline] resume response: {json.dumps(resume_resp, ensure_ascii=False)}")
            final_data = poll_until_terminal(
                session,
                args.base_url.rstrip("/"),
                headers,
                pipeline_id,
                poll_interval=float(args.poll_interval),
                timeout=timeout,
                max_retries=int(args.max_retries),
                deadline_ts=deadline,
                stalled_timeout=float(args.stalled_timeout),
            )
            status = str(final_data.get("pipeline_status") or "")

    print("\n=== FINAL STATUS ===")
    print(json.dumps(final_data, ensure_ascii=False, indent=2))

    if args.save_docx and status in {"completed", "completed_with_warnings"}:
        try:
            out = _save_final_to_docx(
                final_data,
                requested_title=(args.title or "Article"),
                output_path=(args.docx_path or None),
                strip_source_tags=bool(args.strip_source_tags),
            )
            print(f"\n[docx] saved: {out}")
        except Exception as exc:
            print(f"\n[docx] failed to save: {exc}", file=sys.stderr)

    if status in {"completed", "completed_with_warnings"}:
        return 0
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
