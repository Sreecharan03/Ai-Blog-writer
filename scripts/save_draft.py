"""
Download the latest draft for a given request_id from GCS and save as a Word document.
Usage: python scripts/save_draft.py <gcs_uri>
Example: python scripts/save_draft.py gs://ai_blog_02/articles/.../draft_v1/abc123.json
"""
import sys
import json
import re
from pathlib import Path

def main():
    if len(sys.argv) < 2:
        print("Usage: python scripts/save_draft.py <gs://bucket/path/to/draft.json>")
        sys.exit(1)

    gs_uri = sys.argv[1]

    from google.cloud import storage
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Parse GCS URI
    assert gs_uri.startswith("gs://"), "Must be a gs:// URI"
    parts = gs_uri[5:].split("/", 1)
    bucket_name, obj_path = parts[0], parts[1]

    client = storage.Client()
    blob = client.bucket(bucket_name).blob(obj_path)
    raw = blob.download_as_bytes()
    data = json.loads(raw)

    draft = data.get("draft") or {}
    markdown = (draft.get("draft_markdown") or "").strip()
    title = draft.get("title") or "Untitled"
    model = data.get("model", "")
    attempt = data.get("attempt_no", "?")

    if not markdown:
        print("ERROR: draft_markdown is empty in GCS artifact")
        sys.exit(1)

    word_count = len(markdown.split())
    print(f"Draft: {word_count} words | model: {model} | attempt: {attempt}")

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Model: {model} | Attempt: {attempt} | Words: {word_count}")
    doc.add_paragraph("")

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        h3 = re.match(r"^###\s+(.*)", stripped)
        h2 = re.match(r"^##\s+(.*)", stripped)
        h1 = re.match(r"^#\s+(.*)", stripped)

        if h1:
            doc.add_heading(h1.group(1), level=1)
        elif h2:
            doc.add_heading(h2.group(1), level=2)
        elif h3:
            doc.add_heading(h3.group(1), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            # Remove inline markdown bold/italic for plain Word text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            doc.add_paragraph(clean)

    out_path = Path("outputs") / f"draft_{data.get('request_id', 'unknown')[:8]}_attempt{attempt}.docx"
    out_path.parent.mkdir(exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    # Also save raw markdown
    md_path = out_path.with_suffix(".md")
    md_path.write_text(markdown, encoding="utf-8")
    print(f"Saved markdown: {md_path}")

if __name__ == "__main__":
    main()
