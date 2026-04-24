# app/api/url_ingest.py
from __future__ import annotations

import os
import re
import json
import time
import uuid
import html
import hashlib
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import urljoin, urlparse, urlunparse, parse_qsl, urlencode
from urllib import robotparser

from fastapi import APIRouter, BackgroundTasks, Depends, Header, HTTPException, Path
from fastapi.encoders import jsonable_encoder
from pydantic import BaseModel, Field, AnyUrl

import psycopg2
from psycopg2.extras import RealDictCursor
try:
    from google.api_core.exceptions import PreconditionFailed  # type: ignore
except Exception:
    class PreconditionFailed(Exception):
        pass


# ---------------------------
# logging
# ---------------------------
logger = logging.getLogger("url_ingest")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)


# ---------------------------
# Settings loader (supports both patterns)
# ---------------------------
try:
    from app.core.config import get_settings  # type: ignore
except Exception:
    from app.core.config import settings as _settings  # type: ignore

    def get_settings():  # type: ignore
        return _settings


def _pick(settings: Any, *names: str, default: Any = None) -> Any:
    for n in names:
        if hasattr(settings, n) and getattr(settings, n) not in (None, ""):
            return getattr(settings, n)
        n_lower = n.lower()
        if hasattr(settings, n_lower) and getattr(settings, n_lower) not in (None, ""):
            return getattr(settings, n_lower)
        env_val = os.getenv(n)
        if env_val not in (None, ""):
            return env_val
    return default


def _json_dumps_safe(obj: Any) -> str:
    # Ensures Pydantic types (e.g., AnyUrl, datetime) are JSON-serializable.
    return json.dumps(jsonable_encoder(obj))


# ---------------------------
# DB (Supabase Postgres)
# ---------------------------
def _supabase_conn(settings: Any):
    host = _pick(settings, "SUPABASE_DB_HOST", "DB_HOST", "POSTGRES_HOST")
    port = int(_pick(settings, "SUPABASE_DB_PORT", "DB_PORT", "POSTGRES_PORT", default=5432))
    dbname = _pick(settings, "SUPABASE_DB_NAME", "DB_NAME", "POSTGRES_DB", default="postgres")
    user = _pick(settings, "SUPABASE_DB_USER", "DB_USER", "POSTGRES_USER")
    password = _pick(settings, "SUPABASE_DB_PASSWORD", "DB_PASSWORD", "POSTGRES_PASSWORD")
    sslmode = _pick(settings, "SUPABASE_DB_SSLMODE", "DB_SSLMODE", "POSTGRES_SSLMODE", default="require")

    missing = [k for k, v in [("host", host), ("user", user), ("password", password)] if not v]
    if missing:
        raise RuntimeError(f"Missing DB settings: {', '.join(missing)}")

    return psycopg2.connect(
        host=host,
        port=port,
        dbname=dbname,
        user=user,
        password=password,
        sslmode=sslmode,
        connect_timeout=10,
    )


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def _now_iso() -> str:
    return _now_utc().isoformat()


def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _sha256_text(t: str) -> str:
    return hashlib.sha256(t.encode("utf-8", errors="ignore")).hexdigest()


# ---------------------------
# GCS
# ---------------------------
def _gcs_client():
    try:
        from google.cloud import storage  # type: ignore
    except Exception as e:
        raise RuntimeError("google-cloud-storage not installed. pip install google-cloud-storage") from e
    return storage.Client()


def _gcs_upload_bytes(client, bucket_name: str, object_name: str, data: bytes, content_type: str) -> str:
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(object_name)
    try:
        # Create-only (no overwrite => no delete permission needed)
        blob.upload_from_string(data, content_type=content_type, if_generation_match=0)
    except PreconditionFailed:
        # Object already exists => treat as cache hit / reuse
        pass
    return f"gs://{bucket_name}/{object_name}"


# ---------------------------
# Auth (JWT decode)
# ---------------------------
class Claims(BaseModel):
    tenant_id: str
    user_id: str
    role: str
    exp: int


def require_claims(
    authorization: str = Header(..., description="Bearer <JWT>"),
) -> Claims:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing Bearer token")

    token = authorization.split(" ", 1)[1].strip()

    # IMPORTANT: keep consistent with your auth module env
    secret = os.getenv("JWT_SECRET_KEY") or os.getenv("JWT_SECRET") or os.getenv("SECRET_KEY")
    alg = os.getenv("JWT_ALGORITHM") or "HS256"
    if not secret:
        raise HTTPException(status_code=500, detail="JWT secret not configured")

    try:
        from jose import JWTError, jwt  # type: ignore
    except Exception as e:
        raise RuntimeError("python-jose not installed. pip install python-jose[cryptography]") from e

    try:
        payload = jwt.decode(token, secret, algorithms=[alg])
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

    tenant_id = payload.get("tenant_id")
    user_id = payload.get("sub")
    role = payload.get("role")
    exp = payload.get("exp")

    if not tenant_id or not user_id or not role or not exp:
        raise HTTPException(status_code=401, detail="Token missing required claims")

    return Claims(tenant_id=str(tenant_id), user_id=str(user_id), role=str(role), exp=int(exp))


# ---------------------------
# Models
# ---------------------------
class URLIngestRequest(BaseModel):
    url: AnyUrl = Field(..., description="Seed URL to crawl (https://...)")
    max_depth: int = Field(4, ge=0, le=12)
    max_pages: int = Field(300, ge=1, le=5000)
    same_host_only: bool = Field(True, description="Only crawl same host as seed")
    allow_subdomains: bool = Field(False, description="If same_host_only=True, allow subdomains of seed host")
    respect_robots: bool = Field(True)
    user_agent: str = Field("SighnalBot/1.0 (+https://example.invalid/bot)", min_length=3)
    # “auto” = render only if extraction looks too thin (optional)
    render_js: str = Field("auto", description="off|auto|always. Playwright is optional.")
    # guardrails
    max_bytes_per_page: int = Field(8_000_000, ge=50_000, le=30_000_000)
    request_timeout_sec: int = Field(20, ge=3, le=120)
    rate_limit_ms: int = Field(250, ge=0, le=10_000, description="Delay between requests to same host")
    include_patterns: Optional[List[str]] = Field(
        default=None,
        description="Optional regex allowlist applied to normalized URLs",
    )
    exclude_patterns: Optional[List[str]] = Field(
        default=None,
        description="Optional regex blocklist applied to normalized URLs",
    )
    use_default_excludes: bool = Field(
        True,
        description="Apply built-in path excludes like /users/ and /forms/",
    )
    restrict_path_prefix: Optional[str] = Field(
        default=None,
        description="Optional hard path prefix (e.g., /in/en/) to constrain the crawl",
    )
    auto_restrict_locale_prefix: bool = Field(
        True,
        description="If restrict_path_prefix is not set, infer a locale prefix from the seed URL (like /in/en/)",
    )
    wait: bool = Field(False, description="If true, crawl runs inline; else runs in BackgroundTasks")
    auto_pipeline: bool = Field(True, description="Auto preprocess+chunk+embed each page after extraction")


class URLIngestResponse(BaseModel):
    status: str
    kb_id: str
    ingestion_job_id: str
    seed_url: str
    queued_mode: str
    config: Dict[str, Any]


# ---------------------------
# URL normalization & filtering
# ---------------------------
TRACKING_PARAMS_PREFIX = ("utm_",)
TRACKING_PARAMS_EXACT = {"gclid", "fbclid", "yclid", "mc_cid", "mc_eid"}
DEFAULT_EXCLUDE_PATH_PREFIXES = (
    "/users",
    "/user",
    "/account",
    "/profile",
    "/login",
    "/logout",
    "/signup",
    "/register",
    "/forms",
    "/form",
)


def _normalize_url(raw_url: str, base_url: Optional[str] = None) -> Optional[str]:
    try:
        abs_url = urljoin(base_url, raw_url) if base_url else raw_url
        p = urlparse(abs_url)

        if p.scheme not in ("http", "https"):
            return None

        # drop fragments
        fragmentless = p._replace(fragment="")

        # normalize host to lowercase
        netloc = fragmentless.netloc.lower()

        # remove default ports
        if netloc.endswith(":80") and fragmentless.scheme == "http":
            netloc = netloc[:-3]
        if netloc.endswith(":443") and fragmentless.scheme == "https":
            netloc = netloc[:-4]

        # normalize path: collapse multiple slashes
        path = re.sub(r"/{2,}", "/", fragmentless.path or "/")

        # normalize query: drop tracking params, sort
        q = []
        for k, v in parse_qsl(fragmentless.query, keep_blank_values=True):
            lk = k.lower()
            if lk in TRACKING_PARAMS_EXACT:
                continue
            if any(lk.startswith(pref) for pref in TRACKING_PARAMS_PREFIX):
                continue
            q.append((k, v))
        q.sort(key=lambda kv: (kv[0], kv[1]))
        query = urlencode(q, doseq=True)

        normalized = urlunparse((fragmentless.scheme, netloc, path, "", query, ""))
        return normalized
    except Exception:
        return None


def _is_subdomain(host: str, root: str) -> bool:
    host = host.lower()
    root = root.lower()
    return host == root or host.endswith("." + root)


def _matches_any_regex(url: str, patterns: Optional[List[str]]) -> bool:
    if not patterns:
        return False
    for pat in patterns:
        try:
            if re.search(pat, url):
                return True
        except re.error:
            # ignore invalid regex so crawler doesn't crash
            continue
    return False


def _normalize_path_prefix(prefix: str) -> str:
    p = (prefix or "").strip()
    if not p:
        return "/"
    if not p.startswith("/"):
        p = "/" + p
    if not p.endswith("/"):
        p = p + "/"
    return p


def _infer_locale_prefix_from_seed(seed_url: str) -> Optional[str]:
    """
    Infer a locale prefix like /in/en/ from the seed URL path.
    Returns None if no obvious 2-letter/2-letter locale pattern exists.
    """
    try:
        p = urlparse(seed_url)
    except Exception:
        return None

    path = p.path or "/"
    parts = [seg for seg in path.split("/") if seg]
    if len(parts) < 2:
        return None

    a, b = parts[0].lower(), parts[1].lower()
    if re.fullmatch(r"[a-z]{2}", a) and re.fullmatch(r"[a-z]{2}", b):
        return f"/{a}/{b}/"
    return None


def _has_disallowed_prefix(path: str, prefixes: Tuple[str, ...]) -> bool:
    if not prefixes:
        return False
    path_norm = path.lower()
    if not path_norm.endswith("/"):
        path_norm += "/"
    for pref in prefixes:
        pref_norm = _normalize_path_prefix(pref).lower()
        if path_norm.startswith(pref_norm):
            return True
    return False


def _should_keep_url(
    url: str,
    seed_host: str,
    same_host_only: bool,
    allow_subdomains: bool,
    include_patterns: Optional[List[str]],
    exclude_patterns: Optional[List[str]],
    required_path_prefix: Optional[str] = None,
    use_default_excludes: bool = True,
) -> bool:
    p = urlparse(url)
    host = p.netloc.lower()
    path = p.path or "/"

    if same_host_only:
        if allow_subdomains:
            if not _is_subdomain(host, seed_host):
                return False
        else:
            if host != seed_host:
                return False

    if required_path_prefix:
        req = _normalize_path_prefix(required_path_prefix).lower()
        path_cmp = path.lower()
        if not path_cmp.endswith("/"):
            path_cmp += "/"
        if not path_cmp.startswith(req):
            return False

    if use_default_excludes and _has_disallowed_prefix(path, DEFAULT_EXCLUDE_PATH_PREFIXES):
        return False

    # skip obvious non-content
    if re.search(r"\.(jpg|jpeg|png|gif|webp|svg|mp4|mp3|avi|mov|zip|rar|7z)$", p.path, re.I):
        return False

    # apply exclude first
    if _matches_any_regex(url, exclude_patterns):
        return False

    # include allowlist if provided (must match)
    if include_patterns and not _matches_any_regex(url, include_patterns):
        return False

    return True


# ---------------------------
# Robots cache
# ---------------------------
class RobotsCache:
    def __init__(self, user_agent: str, timeout_sec: int):
        self._ua = user_agent
        self._timeout = timeout_sec
        self._cache: Dict[str, robotparser.RobotFileParser] = {}

    def allowed(self, url: str) -> bool:
        p = urlparse(url)
        base = f"{p.scheme}://{p.netloc}"
        if base not in self._cache:
            rp = robotparser.RobotFileParser()
            rp.set_url(base + "/robots.txt")
            try:
                rp.read()
            except Exception:
                # if robots fetch fails, be permissive (configurable at caller)
                pass
            self._cache[base] = rp

        rp = self._cache[base]
        try:
            return rp.can_fetch(self._ua, url)
        except Exception:
            return True


# ---------------------------
# HTTP fetcher (httpx -> requests -> urllib)
# ---------------------------
class FetchResult(BaseModel):
    ok: bool
    status_code: int
    final_url: str
    content_type: str
    content: bytes
    headers: Dict[str, str]


class HttpFetcher:
    def __init__(self, user_agent: str, timeout_sec: int, max_bytes: int):
        self.ua = user_agent
        self.timeout = timeout_sec
        self.max_bytes = max_bytes

        self._httpx = None
        self._requests = None

        try:
            import httpx  # type: ignore

            self._httpx = httpx
        except Exception:
            self._httpx = None

        if self._httpx is None:
            try:
                import requests  # type: ignore

                self._requests = requests
            except Exception:
                self._requests = None

    def get(self, url: str) -> FetchResult:
        headers = {"User-Agent": self.ua, "Accept": "*/*"}

        # 1) httpx
        if self._httpx is not None:
            httpx = self._httpx
            with httpx.Client(follow_redirects=True, timeout=self.timeout, headers=headers) as client:
                r = client.get(url)
                ctype = (r.headers.get("content-type") or "application/octet-stream").split(";")[0].strip().lower()
                content = r.content[: self.max_bytes + 1]
                if len(content) > self.max_bytes:
                    return FetchResult(
                        ok=False,
                        status_code=413,
                        final_url=str(r.url),
                        content_type=ctype,
                        content=b"",
                        headers=dict(r.headers),
                    )
                return FetchResult(
                    ok=200 <= r.status_code < 400,
                    status_code=r.status_code,
                    final_url=str(r.url),
                    content_type=ctype,
                    content=content,
                    headers=dict(r.headers),
                )

        # 2) requests
        if self._requests is not None:
            requests = self._requests
            r = requests.get(url, headers=headers, allow_redirects=True, timeout=self.timeout, stream=True)
            ctype = (r.headers.get("content-type") or "application/octet-stream").split(";")[0].strip().lower()
            buf = bytearray()
            for chunk in r.iter_content(chunk_size=64 * 1024):
                if not chunk:
                    continue
                buf.extend(chunk)
                if len(buf) > self.max_bytes:
                    return FetchResult(
                        ok=False,
                        status_code=413,
                        final_url=r.url,
                        content_type=ctype,
                        content=b"",
                        headers=dict(r.headers),
                    )
            return FetchResult(
                ok=200 <= r.status_code < 400,
                status_code=r.status_code,
                final_url=r.url,
                content_type=ctype,
                content=bytes(buf),
                headers=dict(r.headers),
            )

        # 3) urllib
        import urllib.request

        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=self.timeout) as resp:  # nosec
            status = getattr(resp, "status", 200)
            final_url = resp.geturl()
            ctype = (resp.headers.get("content-type") or "application/octet-stream").split(";")[0].strip().lower()
            content = resp.read(self.max_bytes + 1)
            if len(content) > self.max_bytes:
                return FetchResult(ok=False, status_code=413, final_url=final_url, content_type=ctype, content=b"", headers=dict(resp.headers))
            return FetchResult(ok=200 <= status < 400, status_code=status, final_url=final_url, content_type=ctype, content=content, headers=dict(resp.headers))


# ---------------------------
# Optional JS renderer (Playwright) — used only if installed AND render_js != off
# ---------------------------
def _playwright_available() -> bool:
    try:
        import playwright  # type: ignore  # noqa
        return True
    except Exception:
        return False


def _render_js_if_needed(url: str, user_agent: str, timeout_sec: int) -> Optional[bytes]:
    """
    Best-effort sync Playwright render. If Playwright isn't installed or fails, returns None.
    """
    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except Exception:
        return None

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx = browser.new_context(user_agent=user_agent)
            page = ctx.new_page()
            page.set_default_navigation_timeout(timeout_sec * 1000)
            page.goto(url, wait_until="networkidle")
            # wait for JS stats tables to finish rendering
            page.wait_for_timeout(2500)
            content = page.content().encode("utf-8", errors="ignore")
            browser.close()
            return content
    except Exception:
        return None


# ---------------------------
# Extraction
# ---------------------------
def _decode_text(content: bytes, headers: Dict[str, str]) -> str:
    # try charset in headers
    ct = headers.get("content-type") or headers.get("Content-Type") or ""
    m = re.search(r"charset=([A-Za-z0-9_\-]+)", ct, re.I)
    enc = (m.group(1).strip() if m else "utf-8").lower()

    for candidate in [enc, "utf-8", "utf-16", "latin-1"]:
        try:
            return content.decode(candidate, errors="ignore")
        except Exception:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_links_from_html(html_text: str, base_url: str) -> List[str]:
    links: List[str] = []
    try:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(html_text, "lxml" if _bs4_lxml_available() else "html.parser")
        for a in soup.find_all("a"):
            href = a.get("href")
            if not href:
                continue
            u = _normalize_url(href, base_url=base_url)
            if u:
                links.append(u)
        return links
    except Exception:
        # fallback: regex (best-effort)
        for m in re.finditer(r'href\s*=\s*["\']([^"\']+)["\']', html_text, re.I):
            u = _normalize_url(m.group(1), base_url=base_url)
            if u:
                links.append(u)
        return links


def _bs4_lxml_available() -> bool:
    try:
        import lxml  # type: ignore  # noqa
        return True
    except Exception:
        return False


def _strip_boilerplate(soup) -> None:
    # remove obvious noise
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        tag.decompose()

    # remove nav/footer/header/aside/forms (often cookie/login)
    for tag in soup.find_all(["header", "footer", "nav", "aside", "form"]):
        tag.decompose()

    # heuristic: remove cookie/consent/modals/banners by id/class keywords
    keywords = ("cookie", "consent", "banner", "modal", "popup", "overlay", "subscribe", "newsletter", "otp", "login")
    candidates = soup.find_all(True)
    for el in candidates:
        attrs = " ".join([str(el.get("id", ""))] + [str(x) for x in (el.get("class") or [])]).lower()
        if any(k in attrs for k in keywords):
            try:
                el.decompose()
            except Exception:
                pass


def _pick_main_container(soup):
    # prefer semantic containers
    for sel in ["main", "article"]:
        node = soup.find(sel)
        if node and node.get_text(strip=True):
            return node

    # fallback: choose the node with max text length among div/section
    best = None
    best_len = 0
    for node in soup.find_all(["div", "section"]):
        txt = node.get_text(" ", strip=True)
        if len(txt) > best_len:
            best_len = len(txt)
            best = node
    return best or soup.body or soup


def _normalize_whitespace(text: str) -> str:
    text = html.unescape(text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _universal_post_clean(text: str) -> str:
    """
    Website-agnostic post-extraction cleanup.
    Removes universal noise patterns that appear on ANY website:
    social buttons, copyright lines, duplicate lines, trailing related-content blocks.
    """
    if not text:
        return text
    lines = text.split("\n")
    cleaned: List[str] = []
    prev = ""
    for line in lines:
        s = line.strip()
        if not s:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")  # keep one blank line
            continue
        # Skip exact duplicate consecutive lines
        if s == prev:
            continue
        # Skip near-duplicate consecutive lines (e.g. image caption with/without space before ©)
        if prev and len(s) > 20 and len(prev) > 20:
            # Normalize: strip © and everything after, compare
            s_core = re.sub(r"\s*©.*$", "", s).strip()
            prev_core = re.sub(r"\s*©.*$", "", prev).strip()
            if s_core and s_core == prev_core:
                continue
        # Skip standalone social share buttons
        if s in ("Share", "Tweet", "Pin", "Email", "Print", "Copy Link", "Share Tweet",
                 "Like", "Save", "Bookmark", "Comments", "comment"):
            continue
        # Skip standalone copyright lines (short, starts with ©)
        if s.startswith("\u00a9") and len(s) < 120:
            continue
        # Strip inline © attribution from end of lines (e.g. "Photo caption © AFP")
        s = re.sub(r"\s*©\s*[A-Z][A-Za-z/]*\s*$", "", s).strip()
        if not s:
            continue
        # Skip "by Author •" standalone byline noise (very short with bullet)
        if s == "\u2022" or s == "•":
            continue
        prev = s
        cleaned.append(s)

    # Trim trailing noise: "RELATED STORIES", "LATEST NEWS", "More News", etc.
    # These appear as clusters of short lines (titles of other articles) at the end
    _tail_keywords = {
        "RELATED STORIES", "LATEST NEWS", "More News", "More Stories",
        "READ MORE", "ALSO READ", "TRENDING", "POPULAR", "TOP STORIES",
        "LATEST PHOTOS", "TAGS", "RECOMMENDED", "YOU MAY ALSO LIKE",
        "MOST READ", "EDITOR'S PICKS", "EDITORS PICKS", "SEE ALSO",
    }
    # Walk backwards: remove lines that are part of trailing noise
    while cleaned:
        last = cleaned[-1].strip()
        if not last:
            cleaned.pop()
            continue
        # Exact match on noise headings
        if last.upper() in _tail_keywords:
            cleaned.pop()
            continue
        # Short lines at the very end after a noise heading was already removed
        # (these are article titles from "Related Stories" etc.)
        # Stop if we hit a substantial line (>60 chars or starts with heading marker)
        if len(last) < 80 and not last.startswith("#"):
            # Check if there are more short lines below (cluster detection)
            # We already removed some, so keep going if recent lines are short
            if len(cleaned) >= 2 and len(cleaned[-2].strip()) < 80:
                cleaned.pop()
                continue
        break

    return "\n".join(cleaned).strip()


def _extract_html_trafilatura(html_text: str) -> Optional[str]:
    """Layer 1: trafilatura — ML-trained content extraction. Works on any website."""
    try:
        import trafilatura
        text = trafilatura.extract(
            html_text,
            include_comments=False,
            include_tables=True,
            favor_recall=True,
            deduplicate=True,
        )
        return text if text and len(text) >= 200 else None
    except Exception:
        return None


def _extract_html_readability(html_text: str) -> Optional[str]:
    """Layer 2: readability-lxml — Mozilla Reader Mode algorithm."""
    try:
        from readability import Document
        from bs4 import BeautifulSoup  # type: ignore
        doc = Document(html_text)
        readable_html = doc.summary()
        soup = BeautifulSoup(readable_html, "lxml" if _bs4_lxml_available() else "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text if text and len(text) >= 200 else None
    except Exception:
        return None


def _extract_html_bs4(html_text: str) -> Optional[str]:
    """Layer 3: BeautifulSoup manual extraction — original crawler logic."""
    try:
        from bs4 import BeautifulSoup  # type: ignore
        soup = BeautifulSoup(html_text, "lxml" if _bs4_lxml_available() else "html.parser")
        _strip_boilerplate(soup)
        main = _pick_main_container(soup)

        lines: List[str] = []
        for el in main.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"], recursive=True):
            if el.name in ("h1", "h2", "h3", "h4"):
                t = el.get_text(" ", strip=True)
                if t:
                    lines.append(f"\n# {t}\n")
            elif el.name == "p":
                t = el.get_text(" ", strip=True)
                if t:
                    lines.append(t)
            elif el.name == "li":
                t = el.get_text(" ", strip=True)
                if t:
                    lines.append(f"- {t}")
            elif el.name == "table":
                rows = []
                for tr in el.find_all("tr"):
                    cells = [c.get_text(" ", strip=True) for c in tr.find_all(["th", "td"])]
                    cells = [c for c in cells if c]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    lines.append("\n[Table]\n" + "\n".join(rows) + "\n[/Table]\n")

        text = "\n".join(lines)
        return text if text and len(text) >= 200 else None
    except Exception:
        return None


def _extract_html_meta(html_text: str, url: str) -> Dict[str, Any]:
    """Extract page metadata (title, canonical, description, lang) from HTML."""
    meta: Dict[str, Any] = {"url": url}
    try:
        from bs4 import BeautifulSoup  # type: ignore
        soup = BeautifulSoup(html_text, "lxml" if _bs4_lxml_available() else "html.parser")

        meta["title"] = (soup.title.get_text(strip=True) if soup.title else "").strip()

        canonical = ""
        link_can = soup.find("link", rel=lambda v: v and "canonical" in str(v).lower())
        if link_can and link_can.get("href"):
            canonical = _normalize_url(link_can.get("href"), base_url=url) or ""
        meta["canonical_url"] = canonical

        desc = ""
        mdesc = soup.find("meta", attrs={"name": re.compile(r"^description$", re.I)})
        if mdesc and mdesc.get("content"):
            desc = str(mdesc.get("content")).strip()
        meta["description"] = desc

        lang = ""
        if soup.html and soup.html.get("lang"):
            lang = str(soup.html.get("lang")).strip()
        meta["lang"] = lang
    except Exception:
        pass
    return meta


def extract_from_html(content: bytes, url: str, headers: Dict[str, str]) -> Tuple[str, Dict[str, Any]]:
    """
    3-layer HTML content extraction:
      Layer 1: trafilatura    — ML-trained, best quality, works on any website
      Layer 2: readability    — Mozilla Reader Mode algorithm
      Layer 3: BS4 manual     — original crawler logic (boilerplate strip + main container)
      Layer 4: regex fallback — crude HTML tag stripping

    All layers followed by universal post-clean (dedup, social buttons, copyright, trailing noise).
    """
    html_text = _decode_text(content, headers)
    meta = _extract_html_meta(html_text, url)

    # Layer 1: trafilatura (best quality, ML-based)
    text = _extract_html_trafilatura(html_text)
    if text:
        text = _universal_post_clean(_normalize_whitespace(text))
        if len(text) >= 200:
            meta["extraction_method"] = "html:trafilatura"
            meta["raw_length"] = len(html_text)
            meta["extracted_chars"] = len(text)
            return text, meta

    # Layer 2: readability-lxml (Mozilla Reader Mode)
    text = _extract_html_readability(html_text)
    if text:
        text = _universal_post_clean(_normalize_whitespace(text))
        if len(text) >= 200:
            meta["extraction_method"] = "html:readability"
            meta["raw_length"] = len(html_text)
            meta["extracted_chars"] = len(text)
            return text, meta

    # Layer 3: BS4 manual extraction (original logic)
    text = _extract_html_bs4(html_text)
    if text:
        text = _universal_post_clean(_normalize_whitespace(text))
        if len(text) >= 200:
            meta["extraction_method"] = "html:bs4_main"
            meta["raw_length"] = len(html_text)
            meta["extracted_chars"] = len(text)
            return text, meta

    # Layer 4: regex fallback (last resort)
    crude = re.sub(r"<(script|style).*?>.*?</\1>", " ", html_text, flags=re.I | re.S)
    crude = re.sub(r"<[^>]+>", " ", crude)
    text = _universal_post_clean(_normalize_whitespace(crude))
    meta["extraction_method"] = "html:regex_fallback"
    meta["raw_length"] = len(html_text)
    meta["extracted_chars"] = len(text)
    return text, meta


def extract_from_txt(content: bytes, headers: Dict[str, str]) -> Tuple[str, Dict[str, Any]]:
    txt = _decode_text(content, headers)
    txt = _normalize_whitespace(txt)
    meta = {"extraction_method": "txt:decode", "extracted_chars": len(txt)}
    return txt, meta


def extract_from_docx(content: bytes) -> Tuple[str, Dict[str, Any]]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("python-docx not installed. pip install python-docx") from e

    import io

    f = io.BytesIO(content)
    doc = Document(f)

    lines: List[str] = []
    para_count = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
            para_count += 1

    table_count = 0
    for tbl in doc.tables:
        table_count += 1
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            lines.append("\n[Table]\n" + "\n".join(rows) + "\n[/Table]\n")

    text = _normalize_whitespace("\n".join(lines))
    meta = {
        "extraction_method": "docx:python-docx",
        "extracted_chars": len(text),
        "meta": {"paragraphs": para_count, "tables": table_count},
    }
    return text, meta


def extract_from_pdf(content: bytes) -> Tuple[str, Dict[str, Any]]:
    # best-effort: try pdfminer.six, then pypdf
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text  # type: ignore
        import io

        txt = pdfminer_extract_text(io.BytesIO(content)) or ""
        txt = _normalize_whitespace(txt)
        return txt, {"extraction_method": "pdf:pdfminer.six", "extracted_chars": len(txt)}
    except Exception:
        pass

    try:
        import io
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        txt = _normalize_whitespace("\n".join(parts))
        return txt, {"extraction_method": "pdf:pypdf", "extracted_chars": len(txt)}
    except Exception as e:
        raise RuntimeError(
            "PDF extraction unavailable. Install one:\n"
            "  pip install pdfminer.six\n"
            "or\n"
            "  pip install pypdf"
        ) from e


# ---------------------------
# DB schema bootstrap (safe best-effort)
# ---------------------------
def _ensure_url_pages_table(conn) -> None:
    """
    Creates url_pages table + unique index for idempotency.
    If permissions prevent DDL, we just skip (crawler still works via documents table).
    """
    ddl = """
    CREATE TABLE IF NOT EXISTS public.url_pages (
        page_id uuid PRIMARY KEY,
        tenant_id uuid NOT NULL,
        kb_id uuid NOT NULL,
        doc_id uuid NULL,
        url text NOT NULL,
        normalized_url text NOT NULL,
        parent_url text NULL,
        depth integer NOT NULL DEFAULT 0,
        status_code integer NOT NULL DEFAULT 0,
        content_type text NULL,
        title text NULL,
        canonical_url text NULL,
        raw_fingerprint text NULL,
        text_fingerprint text NULL,
        extracted_chars integer NOT NULL DEFAULT 0,
        discovered_at timestamptz NOT NULL DEFAULT now(),
        crawled_at timestamptz NOT NULL DEFAULT now(),
        meta jsonb NOT NULL DEFAULT '{}'::jsonb
    );
    """
    idx = """
    CREATE UNIQUE INDEX IF NOT EXISTS url_pages_tenant_kb_normurl_ux
    ON public.url_pages (tenant_id, kb_id, normalized_url);
    """
    with conn.cursor() as cur:
        cur.execute(ddl)
        cur.execute(idx)


def _log_job_event(conn, tenant_id: str, request_id: Optional[str], job_id: Optional[str], event_type: str, detail: Optional[Dict[str, Any]] = None):
    detail = detail or {}
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            INSERT INTO public.job_events (tenant_id, request_id, job_id, event_type, detail)
            VALUES (%s, %s, %s, %s, %s::jsonb)
            """,
            (tenant_id, request_id, job_id, event_type, _json_dumps_safe(detail)),
        )


def _insert_document_row(
    conn,
    *,
    doc_id: str,
    tenant_id: str,
    kb_id: str,
    scope: str,
    source_type: str,
    source_name: str,
    mime_type: str,
    byte_size: int,
    fingerprint: str,
    text_fingerprint: Optional[str],
    gcs_raw_uri: str,
    gcs_extracted_uri: Optional[str],
):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            INSERT INTO public.documents
                (doc_id, tenant_id, kb_id, fingerprint, scope, source_type, source_name,
                 mime_type, byte_size, gcs_raw_uri, gcs_extracted_uri, text_fingerprint)
            VALUES
                (%s, %s, %s, %s, %s, %s, %s,
                 %s, %s, %s, %s, %s)
            """,
            (
                doc_id,
                tenant_id,
                kb_id,
                fingerprint,
                scope,
                source_type,
                source_name,
                mime_type,
                byte_size,
                gcs_raw_uri,
                gcs_extracted_uri,
                text_fingerprint,
            ),
        )


def _upsert_url_page(
    conn,
    *,
    tenant_id: str,
    kb_id: str,
    doc_id: Optional[str],
    url: str,
    normalized_url: str,
    parent_url: Optional[str],
    depth: int,
    status_code: int,
    content_type: str,
    title: Optional[str],
    canonical_url: Optional[str],
    raw_fingerprint: Optional[str],
    text_fingerprint: Optional[str],
    extracted_chars: int,
    meta: Dict[str, Any],
):
    page_id = str(uuid.uuid4())
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO public.url_pages
                (page_id, tenant_id, kb_id, doc_id, url, normalized_url, parent_url, depth,
                 status_code, content_type, title, canonical_url, raw_fingerprint, text_fingerprint,
                 extracted_chars, meta)
            VALUES
                (%s, %s, %s, %s, %s, %s, %s, %s,
                 %s, %s, %s, %s, %s, %s,
                 %s, %s::jsonb)
            ON CONFLICT (tenant_id, kb_id, normalized_url)
            DO UPDATE SET
                doc_id = EXCLUDED.doc_id,
                url = EXCLUDED.url,
                parent_url = EXCLUDED.parent_url,
                depth = LEAST(public.url_pages.depth, EXCLUDED.depth),
                status_code = EXCLUDED.status_code,
                content_type = EXCLUDED.content_type,
                title = COALESCE(EXCLUDED.title, public.url_pages.title),
                canonical_url = COALESCE(EXCLUDED.canonical_url, public.url_pages.canonical_url),
                raw_fingerprint = COALESCE(EXCLUDED.raw_fingerprint, public.url_pages.raw_fingerprint),
                text_fingerprint = COALESCE(EXCLUDED.text_fingerprint, public.url_pages.text_fingerprint),
                extracted_chars = GREATEST(public.url_pages.extracted_chars, EXCLUDED.extracted_chars),
                crawled_at = now(),
                meta = public.url_pages.meta || EXCLUDED.meta
            """,
            (
                page_id,
                tenant_id,
                kb_id,
                doc_id,
                url,
                normalized_url,
                parent_url,
                depth,
                status_code,
                content_type,
                title,
                canonical_url,
                raw_fingerprint,
                text_fingerprint,
                extracted_chars,
                _json_dumps_safe(meta),
            ),
        )


# ---------------------------
# Crawl job
# ---------------------------
@dataclass
class CrawlJob:
    tenant_id: str
    user_id: str
    kb_id: str
    job_id: str
    seed_url: str
    seed_host: str
    scope: str
    config: URLIngestRequest


def _content_type_to_kind(content_type: str, url: str) -> str:
    ct = (content_type or "").lower()
    if "text/html" in ct:
        return "html"
    if "application/pdf" in ct or url.lower().endswith(".pdf"):
        return "pdf"
    if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ct or url.lower().endswith(".docx"):
        return "docx"
    if "text/plain" in ct or url.lower().endswith(".txt"):
        return "txt"
    return "bin"


def _safe_object_name(name: str) -> str:
    name = re.sub(r"[^\w\-.() ]+", "_", name)
    name = name.strip().replace(" ", "_")
    return name[:180] if len(name) > 180 else name


def _auto_pipeline_page(
    *,
    doc_id: str,
    tenant_id: str,
    kb_id: str,
    extracted_text: str,
    text_fingerprint: str,
    gcs: Any,
    bucket_name: str,
    settings: Any,
) -> Dict[str, Any]:
    """
    Auto preprocess → chunk → embed a single crawled page.
    Uses core functions directly (no HTTP calls, no FastAPI deps).
    Returns stats dict with pipeline results.
    """
    from app.api.preprocess import clean_text
    from app.api.chunk import chunk_text_dynamic

    pipeline_stats: Dict[str, Any] = {"doc_id": doc_id, "preprocessed": False, "chunked": False, "embedded": False}

    # ── 1. Preprocess (clean text) ──
    try:
        cleaned, clean_stats, clean_meta = clean_text(
            extracted_text,
            remove_boilerplate=True,
            standardize_bullets=True,
            standardize_headings=True,
        )
        if not cleaned or len(cleaned) < 100:
            pipeline_stats["skip_reason"] = "cleaned_text_too_short"
            return pipeline_stats

        clean_fingerprint = _sha256_text(cleaned)
        clean_obj = f"processed/{tenant_id}/{kb_id}/{doc_id}/clean_v1/{clean_fingerprint}.txt"
        gcs_clean_uri = _gcs_upload_bytes(gcs, bucket_name, clean_obj, cleaned.encode("utf-8"), content_type="text/plain; charset=utf-8")

        # Store preprocess output in DB (matches preprocess.py schema exactly)
        with _supabase_conn(settings) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO public.preprocess_outputs
                        (output_id, tenant_id, kb_id, doc_id, preprocessing_version, input_text_fingerprint,
                         clean_fingerprint, gcs_clean_uri, cleaned_chars, method, meta, created_at)
                    VALUES (%s, %s::uuid, %s::uuid, %s::uuid, %s, %s, %s, %s, %s, %s, %s::jsonb, now())
                    ON CONFLICT (tenant_id, kb_id, doc_id, preprocessing_version) DO UPDATE SET
                        input_text_fingerprint = EXCLUDED.input_text_fingerprint,
                        clean_fingerprint = EXCLUDED.clean_fingerprint,
                        gcs_clean_uri = EXCLUDED.gcs_clean_uri,
                        cleaned_chars = EXCLUDED.cleaned_chars,
                        method = EXCLUDED.method,
                        meta = EXCLUDED.meta,
                        created_at = now()
                    """,
                    (str(uuid.uuid4()), tenant_id, kb_id, doc_id, "v1", text_fingerprint,
                     clean_fingerprint, gcs_clean_uri, len(cleaned), "clean:v1:auto_pipeline",
                     json.dumps({"auto_pipeline": True, "input_chars": len(extracted_text), "output_chars": len(cleaned)})),
                )
            conn.commit()
        pipeline_stats["preprocessed"] = True
        pipeline_stats["cleaned_chars"] = len(cleaned)
    except Exception as e:
        logger.warning("auto_pipeline preprocess failed for doc_id=%s: %s", doc_id, str(e)[:200])
        pipeline_stats["preprocess_error"] = str(e)[:200]
        return pipeline_stats

    # ── 2. Chunk ──
    try:
        chunk_size = 2000
        overlap = 200
        max_chunks = 5000
        chunks = chunk_text_dynamic(
            text=cleaned,
            chunk_size_chars=chunk_size,
            overlap_chars=overlap,
            max_chunks=max_chunks,
        )
        if not chunks:
            pipeline_stats["skip_reason"] = "no_chunks"
            return pipeline_stats

        # Build JSONL + manifest
        params_hash = _sha256_text(f"{chunk_size}:{overlap}:{max_chunks}")
        chunk_rows = []
        jsonl_lines = []
        for idx, ch in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            chunk_fp = _sha256_text(ch["text"])
            row = {
                "chunk_id": chunk_id,
                "tenant_id": tenant_id,
                "kb_id": kb_id,
                "doc_id": doc_id,
                "chunk_index": idx,
                "section_path": ch.get("section_path", ""),
                "start_char": ch.get("start_char", 0),
                "end_char": ch.get("end_char", 0),
                "text": ch["text"],
                "chunk_fingerprint": chunk_fp,
            }
            jsonl_lines.append(json.dumps(row, ensure_ascii=False))
            chunk_rows.append(row)

        jsonl_bytes = ("\n".join(jsonl_lines)).encode("utf-8")
        chunks_obj = f"processed/{tenant_id}/{kb_id}/{doc_id}/chunks_v1/{clean_fingerprint}_{params_hash}.jsonl"
        gcs_chunks_uri = _gcs_upload_bytes(gcs, bucket_name, chunks_obj, jsonl_bytes, content_type="application/jsonl; charset=utf-8")

        manifest = {"chunk_count": len(chunks), "params_hash": params_hash, "input_fingerprint": clean_fingerprint}
        manifest_obj = f"processed/{tenant_id}/{kb_id}/{doc_id}/chunks_v1/{clean_fingerprint}_{params_hash}.manifest.json"
        _gcs_upload_bytes(gcs, bucket_name, manifest_obj, json.dumps(manifest).encode("utf-8"), content_type="application/json; charset=utf-8")

        # Store chunks in DB
        with _supabase_conn(settings) as conn:
            with conn.cursor() as cur:
                for row in chunk_rows:
                    cur.execute(
                        """
                        INSERT INTO public.chunks
                            (chunk_id, tenant_id, kb_id, doc_id, chunk_index, section_path,
                             start_char, end_char, chunk_fingerprint, chunk_chars,
                             params_hash, input_fingerprint, input_kind,
                             gcs_chunks_uri, gcs_manifest_uri, created_at)
                        VALUES (%s, %s::uuid, %s::uuid, %s::uuid, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, now())
                        ON CONFLICT (tenant_id, kb_id, doc_id, params_hash, input_fingerprint, chunk_index) DO NOTHING
                        """,
                        (row["chunk_id"], tenant_id, kb_id, doc_id, row["chunk_index"],
                         row["section_path"], row["start_char"], row["end_char"],
                         row["chunk_fingerprint"], len(row["text"]),
                         params_hash, clean_fingerprint, "clean_v1",
                         gcs_chunks_uri, manifest_obj),
                    )
            conn.commit()
        pipeline_stats["chunked"] = True
        pipeline_stats["chunk_count"] = len(chunks)
    except Exception as e:
        logger.warning("auto_pipeline chunk failed for doc_id=%s: %s", doc_id, str(e)[:200])
        pipeline_stats["chunk_error"] = str(e)[:200]
        return pipeline_stats

    # ── 3. Embed ──
    try:
        api_key = None
        for attr in ("GEMINI_API_KEY", "GOOGLE_API_KEY"):
            val = getattr(settings, attr, None) or os.getenv(attr)
            if val:
                api_key = val
                break
        if not api_key:
            pipeline_stats["embed_skipped"] = "no_gemini_api_key"
            return pipeline_stats

        import google.generativeai as genai
        from app.api.embed import _embed_batch
        from psycopg2.extras import execute_values

        genai.configure(api_key=api_key)
        embedding_model = "models/gemini-embedding-001"
        dim = 1536
        batch_size = 32

        chunk_texts = [r["text"] for r in chunk_rows]
        all_vectors: List[List[float]] = []
        for b_start in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[b_start:b_start + batch_size]
            vecs = _embed_batch(model=embedding_model, texts=batch, output_dimensionality=dim)
            all_vectors.extend(vecs)

        if len(all_vectors) != len(chunk_rows):
            pipeline_stats["embed_error"] = f"vector count mismatch: {len(all_vectors)} vs {len(chunk_rows)}"
            return pipeline_stats

        chunks_fp = _sha256_text(clean_fingerprint + params_hash)
        embed_params_hash = _sha256_text(f"{embedding_model}:{dim}")

        def _to_pg_vector(vec: List[float]) -> str:
            return "[" + ",".join(f"{v:.8f}" for v in vec) + "]"

        db_rows = []
        for row, vec in zip(chunk_rows, all_vectors):
            db_rows.append((
                tenant_id, kb_id, doc_id, row["chunk_id"],
                _to_pg_vector(vec), embedding_model, dim,
                chunks_fp, embed_params_hash,
            ))

        with _supabase_conn(settings) as conn:
            with conn.cursor() as cur:
                sql = """
                INSERT INTO public.chunk_embeddings (
                    tenant_id, kb_id, doc_id, chunk_id,
                    embedding, embedding_model, output_dimensionality,
                    chunks_fingerprint, params_hash
                ) VALUES %s
                ON CONFLICT (tenant_id, kb_id, doc_id, chunk_id)
                DO UPDATE SET
                    embedding = EXCLUDED.embedding,
                    embedding_model = EXCLUDED.embedding_model,
                    output_dimensionality = EXCLUDED.output_dimensionality,
                    chunks_fingerprint = EXCLUDED.chunks_fingerprint,
                    params_hash = EXCLUDED.params_hash,
                    created_at = now()
                """
                template = "(%s::uuid,%s::uuid,%s::uuid,%s,%s::vector,%s,%s,%s,%s)"
                execute_values(cur, sql, db_rows, template=template, page_size=200)
            conn.commit()
        pipeline_stats["embedded"] = True
        pipeline_stats["embedded_count"] = len(db_rows)
    except Exception as e:
        logger.warning("auto_pipeline embed failed for doc_id=%s: %s", doc_id, str(e)[:200])
        pipeline_stats["embed_error"] = str(e)[:200]

    return pipeline_stats


def _crawl_and_ingest(job: CrawlJob) -> Dict[str, Any]:
    settings = get_settings()
    bucket_name = _pick(settings, "GCS_BUCKET", "BUCKET_NAME", default=None)
    if not bucket_name:
        raise RuntimeError("Missing GCS_BUCKET (or BUCKET_NAME) in env/settings")

    gcs = _gcs_client()
    fetcher = HttpFetcher(job.config.user_agent, job.config.request_timeout_sec, job.config.max_bytes_per_page)
    robots = RobotsCache(job.config.user_agent, job.config.request_timeout_sec)

    seed_norm = _normalize_url(job.seed_url)
    if not seed_norm:
        raise RuntimeError("Seed URL normalization failed")

    required_path_prefix = None
    if job.config.restrict_path_prefix:
        required_path_prefix = _normalize_path_prefix(job.config.restrict_path_prefix)
    elif job.config.auto_restrict_locale_prefix:
        required_path_prefix = _infer_locale_prefix_from_seed(seed_norm)

    stats = {
        "started_at": _now_iso(),
        "pages_seen": 0,
        "pages_fetched": 0,
        "pages_extracted": 0,
        "pages_skipped": 0,
        "duplicates": 0,
        "robots_blocked": 0,
        "errors": 0,
    }

    with _supabase_conn(settings) as conn:
        # best-effort DDL
        try:
            _ensure_url_pages_table(conn)
        except Exception as e:
            logger.warning("DDL skipped (url_pages): %s", str(e).split("\n")[0])

        _log_job_event(conn, job.tenant_id, None, job.job_id, "crawl_started", {"seed_url": seed_norm, "config": job.config.model_dump()})

    visited: Set[str] = set()
    queue: List[Tuple[str, int, Optional[str]]] = [(seed_norm, 0, None)]  # (url, depth, parent_url)

    last_req_time = 0.0

    while queue and stats["pages_fetched"] < job.config.max_pages:
        url, depth, parent = queue.pop(0)
        stats["pages_seen"] += 1

        if url in visited:
            stats["duplicates"] += 1
            continue
        visited.add(url)

        # scope filter
        if not _should_keep_url(
            url,
            seed_host=job.seed_host,
            same_host_only=job.config.same_host_only,
            allow_subdomains=job.config.allow_subdomains,
            include_patterns=job.config.include_patterns,
            exclude_patterns=job.config.exclude_patterns,
            required_path_prefix=required_path_prefix,
            use_default_excludes=job.config.use_default_excludes,
        ):
            stats["pages_skipped"] += 1
            continue

        # robots
        if job.config.respect_robots and not robots.allowed(url):
            stats["robots_blocked"] += 1
            with _supabase_conn(settings) as conn:
                _log_job_event(conn, job.tenant_id, None, job.job_id, "robots_blocked", {"url": url})
            continue

        # rate-limit (same worker)
        if job.config.rate_limit_ms > 0:
            now = time.time()
            elapsed = now - last_req_time
            wait_s = max(0.0, job.config.rate_limit_ms / 1000.0 - elapsed)
            if wait_s > 0:
                time.sleep(wait_s)
            last_req_time = time.time()

        # fetch
        try:
            fr = fetcher.get(url)
        except Exception as e:
            stats["errors"] += 1
            with _supabase_conn(settings) as conn:
                _log_job_event(conn, job.tenant_id, None, job.job_id, "page_fetch_error", {"url": url, "error": str(e)})
            continue

        stats["pages_fetched"] += 1
        final_url = _normalize_url(fr.final_url) or fr.final_url
        ctype = fr.content_type or "application/octet-stream"

        if fr.status_code == 403:
            stats["pages_skipped"] += 1
            with _supabase_conn(settings) as conn:
                _log_job_event(
                    conn,
                    job.tenant_id,
                    None,
                    job.job_id,
                    "page_blocked",
                    {"url": url, "final_url": final_url, "status_code": fr.status_code},
                )
                try:
                    _upsert_url_page(
                        conn,
                        tenant_id=job.tenant_id,
                        kb_id=job.kb_id,
                        doc_id=None,
                        url=final_url,
                        normalized_url=final_url,
                        parent_url=parent,
                        depth=depth,
                        status_code=fr.status_code,
                        content_type=ctype,
                        title=None,
                        canonical_url=None,
                        raw_fingerprint=None,
                        text_fingerprint=None,
                        extracted_chars=0,
                        meta={"blocked": True},
                    )
                except Exception:
                    pass
            continue

        # oversize
        if fr.status_code == 413:
            with _supabase_conn(settings) as conn:
                _log_job_event(conn, job.tenant_id, None, job.job_id, "page_skipped_oversize", {"url": url, "final_url": final_url})
            stats["pages_skipped"] += 1
            continue

        # store raw
        raw_fingerprint = _sha256_bytes(fr.content) if fr.content else None
        url_key = _sha256_text(final_url)[:24]
        raw_suffix = raw_fingerprint or "raw"
        raw_obj = f"url_snapshots/{job.tenant_id}/{job.kb_id}/{job.job_id}/{url_key}/raw_{raw_suffix}.raw"
        raw_uri = ""
        try:
            raw_uri = _gcs_upload_bytes(gcs, bucket_name, raw_obj, fr.content, content_type=ctype)
        except Exception as e:
            stats["errors"] += 1
            with _supabase_conn(settings) as conn:
                _log_job_event(conn, job.tenant_id, None, job.job_id, "gcs_upload_error", {"url": final_url, "error": str(e)})
            continue

        # extract
        extracted_text = ""
        extraction_ok = False
        extract_meta: Dict[str, Any] = {}
        extract_method = "none"
        kind = _content_type_to_kind(ctype, final_url)

        try:
            if kind == "html":
                extracted_text, extract_meta = extract_from_html(fr.content, final_url, fr.headers)
                extract_method = extract_meta.get("extraction_method", "html")
                extraction_ok = len(extracted_text.split()) >= 150  # word-count threshold — catches JS pages that return nav noise but no real content
                # JS auto-render if too thin and allowed
                if (not extraction_ok) and job.config.render_js in ("auto", "always") and _playwright_available():
                    rendered = _render_js_if_needed(final_url, job.config.user_agent, job.config.request_timeout_sec)
                    if rendered:
                        extracted_text2, extract_meta2 = extract_from_html(rendered, final_url, {"content-type": "text/html; charset=utf-8"})
                        if len(extracted_text2) > len(extracted_text):
                            extracted_text, extract_meta = extracted_text2, extract_meta2
                            extract_method = extract_meta.get("extraction_method", "html") + "+playwright"
                            extraction_ok = len(extracted_text) >= 300
            elif kind == "txt":
                extracted_text, extract_meta = extract_from_txt(fr.content, fr.headers)
                extract_method = extract_meta.get("extraction_method", "txt")
                extraction_ok = len(extracted_text) >= 50
            elif kind == "docx":
                extracted_text, extract_meta = extract_from_docx(fr.content)
                extract_method = extract_meta.get("extraction_method", "docx")
                extraction_ok = len(extracted_text) >= 100
            elif kind == "pdf":
                extracted_text, extract_meta = extract_from_pdf(fr.content)
                extract_method = extract_meta.get("extraction_method", "pdf")
                extraction_ok = len(extracted_text) >= 200
            else:
                extraction_ok = False
                extract_method = f"unsupported:{kind}"
        except Exception as e:
            extraction_ok = False
            extract_method = f"error:{kind}"
            extract_meta = {"error": str(e)}

        text_fingerprint = _sha256_text(extracted_text) if extracted_text else None
        extracted_chars = len(extracted_text) if extracted_text else 0

        extracted_uri = None
        if extraction_ok and extracted_text:
            extracted_obj = f"processed/{job.tenant_id}/{job.kb_id}/{job.job_id}/{url_key}/extracted_{text_fingerprint}.txt"
            try:
                extracted_uri = _gcs_upload_bytes(gcs, bucket_name, extracted_obj, extracted_text.encode("utf-8"), content_type="text/plain; charset=utf-8")
            except Exception as e:
                with _supabase_conn(settings) as conn:
                    _log_job_event(conn, job.tenant_id, None, job.job_id, "gcs_upload_error", {"url": final_url, "error": str(e)})
                extracted_uri = None

        # insert DB rows
        doc_id = str(uuid.uuid4())
        source_name = extract_meta.get("title") or final_url
        source_name = _safe_object_name(str(source_name)) or final_url

        with _supabase_conn(settings) as conn:
            # job event: fetched
            _log_job_event(conn, job.tenant_id, None, job.job_id, "page_fetched", {"url": url, "final_url": final_url, "status_code": fr.status_code, "content_type": ctype, "depth": depth})

            # url_pages upsert (hierarchy tracking) - works even if documents insert fails
            try:
                _upsert_url_page(
                    conn,
                    tenant_id=job.tenant_id,
                    kb_id=job.kb_id,
                    doc_id=None,  # filled after documents insert
                    url=final_url,
                    normalized_url=final_url,
                    parent_url=parent,
                    depth=depth,
                    status_code=fr.status_code,
                    content_type=ctype,
                    title=extract_meta.get("title"),
                    canonical_url=extract_meta.get("canonical_url"),
                    raw_fingerprint=raw_fingerprint,
                    text_fingerprint=text_fingerprint,
                    extracted_chars=extracted_chars,
                    meta={
                        "seed_url": seed_norm,
                        "extraction_ok": extraction_ok,
                        "extraction_method": extract_method,
                        "lang": extract_meta.get("lang"),
                        "description": extract_meta.get("description"),
                    },
                )
            except Exception:
                # url_pages table might not exist / no privileges
                pass

            if extraction_ok:
                # Content-level dedup: skip if identical text already exists in this KB
                if text_fingerprint:
                    try:
                        with conn.cursor() as _cur:
                            _cur.execute(
                                "SELECT doc_id FROM public.documents WHERE tenant_id=%s::uuid AND kb_id=%s::uuid AND text_fingerprint=%s LIMIT 1",
                                (job.tenant_id, job.kb_id, text_fingerprint),
                            )
                            _existing = _cur.fetchone()
                        if _existing:
                            _log_job_event(conn, job.tenant_id, None, job.job_id, "page_skipped_duplicate",
                                           {"url": final_url, "existing_doc_id": str(_existing[0]), "text_fingerprint": text_fingerprint})
                            conn.commit()
                            continue
                    except Exception:
                        pass  # dedup check failed — proceed with normal insert

                try:
                    _insert_document_row(
                        conn,
                        doc_id=doc_id,
                        tenant_id=job.tenant_id,
                        kb_id=job.kb_id,
                        scope=job.scope,
                        source_type="url",
                        source_name=source_name,
                        mime_type=ctype,
                        byte_size=len(fr.content),
                        fingerprint=raw_fingerprint or _sha256_text(final_url),
                        text_fingerprint=text_fingerprint,
                        gcs_raw_uri=raw_uri,
                        gcs_extracted_uri=extracted_uri,
                    )
                    # update url_pages with doc_id (best-effort)
                    try:
                        _upsert_url_page(
                            conn,
                            tenant_id=job.tenant_id,
                            kb_id=job.kb_id,
                            doc_id=doc_id,
                            url=final_url,
                            normalized_url=final_url,
                            parent_url=parent,
                            depth=depth,
                            status_code=fr.status_code,
                            content_type=ctype,
                            title=extract_meta.get("title"),
                            canonical_url=extract_meta.get("canonical_url"),
                            raw_fingerprint=raw_fingerprint,
                            text_fingerprint=text_fingerprint,
                            extracted_chars=extracted_chars,
                            meta={"document_inserted": True},
                        )
                    except Exception:
                        pass

                    _log_job_event(conn, job.tenant_id, None, job.job_id, "page_extracted", {"url": final_url, "doc_id": doc_id, "extracted_chars": extracted_chars, "method": extract_method})
                    stats["pages_extracted"] += 1

                    # ── Auto-pipeline: preprocess → chunk → embed ──
                    if job.config.auto_pipeline and extracted_text:
                        try:
                            pipe_result = _auto_pipeline_page(
                                doc_id=doc_id,
                                tenant_id=job.tenant_id,
                                kb_id=job.kb_id,
                                extracted_text=extracted_text,
                                text_fingerprint=text_fingerprint or "",
                                gcs=gcs,
                                bucket_name=bucket_name,
                                settings=settings,
                            )
                            _log_job_event(conn, job.tenant_id, None, job.job_id, "auto_pipeline_done", {"doc_id": doc_id, "result": pipe_result})
                            stats.setdefault("pages_pipelined", 0)
                            if pipe_result.get("embedded"):
                                stats["pages_pipelined"] += 1
                        except Exception as e:
                            logger.warning("auto_pipeline error doc_id=%s: %s", doc_id, str(e)[:200])

                except psycopg2.Error as e:
                    _log_job_event(conn, job.tenant_id, None, job.job_id, "db_insert_error", {"url": final_url, "error": str(e).split("\n")[0]})
                    stats["errors"] += 1
            else:
                _log_job_event(conn, job.tenant_id, None, job.job_id, "page_not_extracted", {"url": final_url, "reason": extract_method, "extracted_chars": extracted_chars})

        # expand links
        if kind == "html" and depth < job.config.max_depth and fr.ok and fr.content:
            try:
                html_text = _decode_text(fr.content, fr.headers)
                links = _extract_links_from_html(html_text, base_url=final_url)
                for link in links:
                    if link in visited:
                        continue
                    if not _should_keep_url(
                        link,
                        seed_host=job.seed_host,
                        same_host_only=job.config.same_host_only,
                        allow_subdomains=job.config.allow_subdomains,
                        include_patterns=job.config.include_patterns,
                        exclude_patterns=job.config.exclude_patterns,
                        required_path_prefix=required_path_prefix,
                        use_default_excludes=job.config.use_default_excludes,
                    ):
                        continue
                    queue.append((link, depth + 1, final_url))
            except Exception:
                pass

    stats["finished_at"] = _now_iso()

    with _supabase_conn(settings) as conn:
        _log_job_event(conn, job.tenant_id, None, job.job_id, "crawl_completed", {"stats": stats})

    return stats


# ---------------------------
# Router
# ---------------------------
router = APIRouter(prefix="/api/v1/kb", tags=["ingest-url"])


@router.post("/{kb_id}/ingest/url", response_model=URLIngestResponse)
def ingest_url(
    payload: URLIngestRequest,
    background: BackgroundTasks,
    kb_id: str = Path(..., description="KB UUID"),
    claims: Claims = Depends(require_claims),
):
    # validate kb uuid
    try:
        _ = uuid.UUID(kb_id)
    except Exception:
        raise HTTPException(status_code=400, detail="kb_id must be a valid UUID")

    seed_url = str(payload.url)
    seed_norm = _normalize_url(seed_url)
    if not seed_norm:
        raise HTTPException(status_code=400, detail="Invalid URL")

    seed_host = urlparse(seed_norm).netloc.lower()
    if payload.restrict_path_prefix:
        required = _normalize_path_prefix(payload.restrict_path_prefix).lower()
        seed_path = (urlparse(seed_norm).path or "/").lower()
        if not seed_path.endswith("/"):
            seed_path += "/"
        if not seed_path.startswith(required):
            raise HTTPException(
                status_code=400,
                detail="Seed URL path is outside restrict_path_prefix",
            )
    job_id = str(uuid.uuid4())

    # scope based on your plan (tenant_private by default)
    scope = "tenant_private"

    job = CrawlJob(
        tenant_id=claims.tenant_id,
        user_id=claims.user_id,
        kb_id=kb_id,
        job_id=job_id,
        seed_url=seed_norm,
        seed_host=seed_host,
        scope=scope,
        config=payload,
    )

    # inline vs background
    if payload.wait:
        stats = _crawl_and_ingest(job)
        return URLIngestResponse(
            status="ok",
            kb_id=kb_id,
            ingestion_job_id=job_id,
            seed_url=seed_norm,
            queued_mode="inline",
            config={"stats": stats, **payload.model_dump()},
        )

    background.add_task(_crawl_and_ingest, job)
    return URLIngestResponse(
        status="ok",
        kb_id=kb_id,
        ingestion_job_id=job_id,
        seed_url=seed_norm,
        queued_mode="background",
        config=payload.model_dump(),
    )
