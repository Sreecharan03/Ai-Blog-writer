from __future__ import annotations

import hashlib
import json
import os
import re
import unicodedata
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, Optional, Tuple

from fastapi import APIRouter, Depends, File, Header, HTTPException, UploadFile
from jose import JWTError, jwt
from pydantic import BaseModel, Field

import psycopg2
from psycopg2.extras import Json, RealDictCursor

from google.cloud import storage
from google.api_core.exceptions import PreconditionFailed

# Optional libs (installed in your env already)
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pypdf import PdfReader
from docx import Document as DocxDocument
from fastapi.encoders import jsonable_encoder
from app.services.cache_registry import CacheRegistryService, CacheScope

try:
    from app.core.config import Config  # type: ignore
except Exception:
    class Config:  # type: ignore
        PREPROCESSING_VERSION = os.getenv("PREPROCESSING_VERSION", "v1")

# Keep this near other singletons/config
PREPROCESSING_VERSION = getattr(Config, "PREPROCESSING_VERSION", "v1")

_cache = CacheRegistryService(preprocessing_version=PREPROCESSING_VERSION)
_cache_ready = False


def _get_cache() -> CacheRegistryService:
    global _cache_ready
    if not _cache_ready:
        # Safe to call multiple times; it is idempotent in our cache_registry.py
        _cache.ensure_schema()
        _cache_ready = True
    return _cache
# ---------------------------
# Settings loader (your existing pattern)
# ---------------------------
try:
    from app.core.config import get_settings  # type: ignore
except Exception:
    # fallback: app.core.config.settings
    from app.core.config import settings as _settings  # type: ignore

    def get_settings():  # type: ignore
        return _settings


def _pick(settings: Any, *names: str, default: Any = None) -> Any:
    for n in names:
        if hasattr(settings, n) and getattr(settings, n) not in (None, ""):
            return getattr(settings, n)
        n_lower = n.lower()
        if hasattr(settings, n_lower) and getattr(settings, n_lower) not in (None, ""):
            return getattr(settings, n_lower)
        env_val = os.getenv(n)
        if env_val not in (None, ""):
            return env_val
    return default


# ---------------------------
# JWT Claims (same style as finance)
# ---------------------------
class Claims(BaseModel):
    tenant_id: str
    user_id: str
    role: str
    exp: int


def require_claims(
    authorization: str = Header(..., description="Bearer <JWT>"),
) -> Claims:
    if not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing Bearer token")

    token = authorization.split(" ", 1)[1].strip()

    secret = os.getenv("JWT_SECRET_KEY") or os.getenv("JWT_SECRET") or os.getenv("SECRET_KEY")
    alg = os.getenv("JWT_ALGORITHM") or "HS256"
    if not secret:
        raise HTTPException(status_code=500, detail="JWT secret not configured")

    try:
        payload = jwt.decode(token, secret, algorithms=[alg])
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

    tenant_id = payload.get("tenant_id")
    user_id = payload.get("sub")
    role = payload.get("role")
    exp = payload.get("exp")

    if not tenant_id or not user_id or not role or not exp:
        raise HTTPException(status_code=401, detail="Token missing required claims")

    return Claims(tenant_id=str(tenant_id), user_id=str(user_id), role=str(role), exp=int(exp))


# ---------------------------
# DB connection (uses your current Supabase Postgres pooler envs)
# NOTE: later you can point these to Cloud SQL and it will still work.
# ---------------------------
def _db_conn(settings: Any):
    host = _pick(settings, "DB_HOST", "SUPABASE_DB_HOST", "POSTGRES_HOST")
    port = int(_pick(settings, "DB_PORT", "SUPABASE_DB_PORT", "POSTGRES_PORT", default=5432))
    dbname = _pick(settings, "DB_NAME", "SUPABASE_DB_NAME", "POSTGRES_DB", default="postgres")
    user = _pick(settings, "DB_USER", "SUPABASE_DB_USER", "POSTGRES_USER")
    password = _pick(settings, "DB_PASSWORD", "SUPABASE_DB_PASSWORD", "POSTGRES_PASSWORD")
    sslmode = _pick(settings, "DB_SSLMODE", "SUPABASE_DB_SSLMODE", "POSTGRES_SSLMODE", default="require")

    missing = [k for k, v in [("host", host), ("user", user), ("password", password)] if not v]
    if missing:
        raise RuntimeError(f"Missing DB settings: {', '.join(missing)}")

    return psycopg2.connect(
        host=host,
        port=port,
        dbname=dbname,
        user=user,
        password=password,
        sslmode=sslmode,
        connect_timeout=8,
    )


def _ensure_documents_table(conn) -> None:
    """
    Dev-safe bootstrap so ingestion won't fail on a fresh DB.
    Ensures columns needed by /kb/{kb_id}/docs are present (including extraction_ok).
    """
    with conn.cursor() as cur:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS public.documents (
              doc_id uuid PRIMARY KEY,
              tenant_id uuid NOT NULL,
              kb_id uuid NOT NULL,
              fingerprint text NOT NULL,
              scope text NOT NULL DEFAULT 'tenant_private',
              source_type text NOT NULL,
              source_name text,
              mime_type text,
              byte_size bigint,
              gcs_raw_uri text,
              gcs_extracted_uri text,
              text_fingerprint text,
              extraction_ok boolean,
              extraction_method text,
              extracted_chars integer,
              created_at timestamptz NOT NULL DEFAULT now()
            );
            """
        )
        # Backward-compatible adds (if table existed before)
        cur.execute("ALTER TABLE public.documents ADD COLUMN IF NOT EXISTS extraction_ok boolean;")
        cur.execute("ALTER TABLE public.documents ADD COLUMN IF NOT EXISTS extraction_method text;")
        cur.execute("ALTER TABLE public.documents ADD COLUMN IF NOT EXISTS extracted_chars integer;")
        cur.execute(
            "CREATE INDEX IF NOT EXISTS idx_documents_tenant_kb_created_at ON public.documents (tenant_id, kb_id, created_at DESC);"
        )
        cur.execute("CREATE INDEX IF NOT EXISTS idx_documents_fingerprint ON public.documents (fingerprint);")


def _log_job_event(conn, tenant_id: str, event_type: str, detail: Dict[str, Any], job_id: Optional[str] = None) -> None:
    """
    Writes into your existing Supabase job_events table (already created earlier).
    """
    # If job_events table does not exist, fail loudly (we don't want silent logging loss).
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO public.job_events (event_id, tenant_id, request_id, job_id, event_type, detail)
            VALUES (%s, %s, %s, %s, %s, %s)
            """,
            (
                str(uuid.uuid4()),
                tenant_id,
                None,
                job_id,
                event_type,
                Json(detail),
            ),
        )


# ---------------------------
# GCS helpers
# ---------------------------
def _norm_prefix(p: str) -> str:
    p = (p or "").strip().replace("\\", "/")
    if not p:
        return ""
    return p if p.endswith("/") else (p + "/")


def _safe_name(name: str) -> str:
    name = name.strip().replace("\\", "/").split("/")[-1]
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
    return name[:160] if len(name) > 160 else name


def _gcs_client(settings: Any) -> storage.Client:
    project_id = _pick(settings, "GCP_PROJECT_ID", "GOOGLE_CLOUD_PROJECT")
    return storage.Client(project=project_id)


def _gcs_upload_bytes(
    client: storage.Client,
    bucket_name: str,
    object_name: str,
    data: bytes,
    content_type: str,
) -> str:
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(object_name)
    try:
        # Create-only (no overwrite => no delete permission needed)
        blob.upload_from_string(data, content_type=content_type, if_generation_match=0)
    except PreconditionFailed:
        # Object already exists => treat as cache hit / reuse
        pass
    return f"gs://{bucket_name}/{object_name}"


# ---------------------------
# Extraction (best-effort, no OCR)
# ---------------------------
@dataclass
class ExtractResult:
    ok: bool
    method: str
    text: str
    meta: Dict[str, Any]


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _normalize_text(t: str) -> str:
    t = unicodedata.normalize("NFKC", t)
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # very light cleanup; heavy cleanup is Day-11 preprocess
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t).strip()
    return t


def _extract_txt(data: bytes) -> ExtractResult:
    # try utf-8-sig first, then utf-8, then latin1
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            txt = data.decode(enc)
            return ExtractResult(True, f"txt:{enc}", _normalize_text(txt), {"encoding": enc})
        except Exception:
            continue
    txt = data.decode("utf-8", errors="replace")
    return ExtractResult(True, "txt:replace", _normalize_text(txt), {"encoding": "utf-8:replace"})


def _extract_docx(data: bytes) -> ExtractResult:
    doc = DocxDocument(BytesIO(data))
    parts = []
    # paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # tables
    for table in doc.tables:
        for row in table.rows:
            row_txt = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_txt:
                parts.append(" | ".join(row_txt))
    text = "\n".join(parts)
    text = _normalize_text(text)
    return ExtractResult(True, "docx:python-docx", text, {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)})


def _extract_pdf(data: bytes) -> ExtractResult:
    meta: Dict[str, Any] = {}
    # page count (fast)
    try:
        reader = PdfReader(BytesIO(data))
        meta["pages"] = len(reader.pages)
    except Exception:
        meta["pages"] = None

    # primary: pdfminer
    text_a = ""
    try:
        text_a = pdfminer_extract_text(BytesIO(data)) or ""
    except Exception as e:
        meta["pdfminer_error"] = str(e).split("\n")[0]

    text_a_n = _normalize_text(text_a) if text_a else ""

    # fallback: pypdf extraction if pdfminer is weak
    text_b = ""
    try:
        reader = PdfReader(BytesIO(data))
        out = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                out.append(t)
        text_b = "\n\n".join(out)
    except Exception as e:
        meta["pypdf_error"] = str(e).split("\n")[0]

    text_b_n = _normalize_text(text_b) if text_b else ""

    # choose better
    if len(text_b_n) > max(200, len(text_a_n) * 2):
        return ExtractResult(True, "pdf:pypdf", text_b_n, meta)

    if len(text_a_n) > 0:
        return ExtractResult(True, "pdf:pdfminer", text_a_n, meta)

    return ExtractResult(False, "pdf:none", "", meta)


def _extract_by_filename(filename: str, data: bytes) -> ExtractResult:
    ext = (filename.split(".")[-1] if "." in filename else "").lower()

    if ext in ("txt", "md", "csv", "log"):
        return _extract_txt(data)

    if ext in ("docx",):
        return _extract_docx(data)

    if ext in ("pdf",):
        return _extract_pdf(data)

    # accept .doc but no extractor installed (no mock OCR)
    if ext in ("doc",):
        return ExtractResult(False, "doc:unsupported", "", {"note": "Upload stored; extraction requires .docx or server-side converter."})

    return ExtractResult(False, f"{ext or 'unknown'}:unsupported", "", {"note": "Unsupported file type for extraction."})


# ---------------------------
# API Router (Day-7 Ingestion)
# ---------------------------
router = APIRouter(prefix="/api/v1/kb", tags=["ingest"])


class IngestResponse(BaseModel):
    status: str
    kb_id: str
    doc_id: str
    ingestion_job_id: str
    fingerprint: str
    text_fingerprint: Optional[str] = None
    scope: str
    gcs_raw_uri: str
    gcs_extracted_uri: Optional[str] = None
    extraction_ok: bool
    extraction_method: str
    extracted_chars: int
    meta: Dict[str, Any] = Field(default_factory=dict)


@router.post("/{kb_id}/ingest/file", response_model=IngestResponse)
async def ingest_file(
    kb_id: str,
    file: UploadFile = File(...),
    claims: Claims = Depends(require_claims),
):
    """
    Day-7: File ingestion endpoint
    upload → GCS raw → fingerprint → (cache.lookup BEFORE extraction)
    If cache hit: skip extraction + skip processed write; reuse gcs_extracted_uri + text_fingerprint
    Else: extract → GCS processed → cache.upsert → documents row → job_events.
    """
    # validate kb_id UUID
    try:
        _ = uuid.UUID(kb_id)
    except Exception:
        raise HTTPException(status_code=400, detail="kb_id must be a valid UUID")

    settings = get_settings()

    # limits (default 25MB)
    max_mb = float(_pick(settings, "MAX_UPLOAD_MB", default=25))
    max_bytes = int(max_mb * 1024 * 1024)

    original_name = _safe_name(file.filename or "upload.bin")
    content_type = file.content_type or "application/octet-stream"

    ingestion_job_id = str(uuid.uuid4())
    tenant_id = claims.tenant_id

    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="Empty upload")
    if len(raw_bytes) > max_bytes:
        raise HTTPException(status_code=413, detail=f"File too large. Max {max_mb} MB")

    fingerprint = _sha256_bytes(raw_bytes)

    # Build GCS paths
    bucket_name = _pick(settings, "GCS_BUCKET_NAME")
    if not bucket_name:
        raise HTTPException(status_code=500, detail="Missing GCS_BUCKET_NAME")

    raw_prefix = _norm_prefix(_pick(settings, "GCS_RAW_PREFIX", default="raw"))
    processed_prefix = _norm_prefix(_pick(settings, "GCS_PROCESSED_PREFIX", default="processed"))

    # doc_id created now (logical doc row)
    doc_id = str(uuid.uuid4())

    # Store raw
    client = _gcs_client(settings)
    ext = os.path.splitext(original_name)[1].lower() or ".bin"
    raw_object = f"{raw_prefix}{tenant_id}/{kb_id}/{doc_id}/raw/{fingerprint}{ext}"
    gcs_raw_uri = _gcs_upload_bytes(client, bucket_name, raw_object, raw_bytes, content_type=content_type)

    # ---- Day-9 cache lookup BEFORE extraction ----
    cache = _get_cache()

    # Map KB/doc scope → cache scope (KBs are tenant_private by default)
    kb_scope_str = "tenant_private"
    if kb_scope_str in ("global_public", "public"):
        desired_scope = CacheScope.GLOBAL_PUBLIC
        allow_global_fallback = True
    else:
        desired_scope = CacheScope.TENANT_PRIVATE
        allow_global_fallback = False

    cache_hit = cache.lookup(
        fingerprint=fingerprint,
        tenant_id=tenant_id,
        desired_scope=desired_scope,
        allow_global_fallback=allow_global_fallback,
        require_extracted_artifact=True,
    )

    gcs_extracted_uri: Optional[str] = None
    text_fingerprint: Optional[str] = None
    extraction_ok = False
    extraction_method = "none"
    extracted_chars = 0
    extract_meta: Dict[str, Any] = {}
    extracted_text: Optional[str] = None
    cache_hit_event: Optional[Dict[str, Any]] = None
    cache_upsert_event: Optional[Dict[str, Any]] = None

    # Treat as hit ONLY if required artifacts exist (extra safety)
    if cache_hit and cache_hit.gcs_extracted_uri and cache_hit.text_fingerprint:
        extraction_ok = True
        gcs_extracted_uri = cache_hit.gcs_extracted_uri
        text_fingerprint = cache_hit.text_fingerprint
        prior_method = cache_hit.extraction_method or "unknown"
        extraction_method = f"cache_hit:{prior_method}"
        extracted_chars = int((cache_hit.meta or {}).get("extracted_chars", 0))
        extract_meta = cache_hit.meta or {}

        cache_hit_event = jsonable_encoder(
            {
                "fingerprint": fingerprint,
                "cache_id": cache_hit.cache_id,
                "cache_scope": cache_hit.scope,
                "gcs_extracted_uri": gcs_extracted_uri,
            }
        )
    else:
        # Cache miss → Extract best-effort
        extract_res = _extract_by_filename(original_name, raw_bytes)
        extraction_ok = bool(extract_res.ok)
        extraction_method = extract_res.method
        extract_meta = extract_res.meta
        extracted_text = extract_res.text
        extracted_chars = len(extract_res.text or "")

        if extraction_ok and extracted_text and extracted_text.strip():
            text_bytes = extracted_text.encode("utf-8")
            text_fingerprint = _sha256_bytes(text_bytes)
            extracted_object = f"{processed_prefix}{tenant_id}/{kb_id}/{doc_id}/extracted/{text_fingerprint}.txt"
            gcs_extracted_uri = _gcs_upload_bytes(
                client, bucket_name, extracted_object, text_bytes, content_type="text/plain; charset=utf-8"
            )

        # Cache upsert AFTER we have a real extracted artifact
        if extraction_ok and gcs_extracted_uri and text_fingerprint:
            cache.upsert(
                fingerprint=fingerprint,
                text_fingerprint=text_fingerprint,
                scope=desired_scope,
                tenant_id=tenant_id,
                canonical_url=None,
                gcs_raw_uri=gcs_raw_uri,
                gcs_extracted_uri=gcs_extracted_uri,
                extraction_method=extraction_method,
                meta={
                    "source_type": "file",
                    "source_name": original_name,
                    "mime_type": content_type,
                    "byte_size": len(raw_bytes),
                    "extracted_chars": int(extracted_chars),
                },
            )

            cache_upsert_event = jsonable_encoder(
                {
                    "fingerprint": fingerprint,
                    "scope": desired_scope.value,
                    "gcs_extracted_uri": gcs_extracted_uri,
                    "text_fingerprint": text_fingerprint,
                }
            )

    # Persist document row + job events (real DB)
    with _db_conn(settings) as conn:
        _ensure_documents_table(conn)

        # job events: started + raw uploaded + extracted (if any) + doc row created + done
        _log_job_event(conn, tenant_id, "step_started", {"step": "ingest_file", "kb_id": kb_id, "filename": original_name}, job_id=ingestion_job_id)
        _log_job_event(conn, tenant_id, "raw_uploaded", {"gcs_raw_uri": gcs_raw_uri, "bytes": len(raw_bytes)}, job_id=ingestion_job_id)

        if cache_hit_event:
            _log_job_event(conn, tenant_id, "cache_hit", cache_hit_event, job_id=ingestion_job_id)

        if cache_upsert_event:
            _log_job_event(
                conn,
                tenant_id,
                "cache_registry_upserted",
                cache_upsert_event,
                job_id=ingestion_job_id,
            )

        if gcs_extracted_uri and not cache_hit_event:
            _log_job_event(
                conn,
                tenant_id,
                "extracted_text_saved",
                {"gcs_extracted_uri": gcs_extracted_uri, "method": extraction_method, "chars": extracted_chars},
                job_id=ingestion_job_id,
            )
        elif not cache_hit_event:
            _log_job_event(
                conn,
                tenant_id,
                "extraction_skipped",
                {"method": extraction_method, "ok": extraction_ok, "note": extract_meta.get("note")},
                job_id=ingestion_job_id,
            )

        # Insert documents row (NOW includes extraction_ok fields)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO public.documents
                  (doc_id, tenant_id, kb_id, fingerprint, scope, source_type, source_name, mime_type, byte_size,
                   gcs_raw_uri, gcs_extracted_uri, text_fingerprint, extraction_ok, extraction_method, extracted_chars)
                VALUES
                  (%s::uuid, %s::uuid, %s::uuid, %s, %s, %s, %s, %s, %s,
                   %s, %s, %s, %s, %s, %s)
                """,
                (
                    doc_id,
                    tenant_id,
                    kb_id,
                    fingerprint,
                    kb_scope_str,
                    "file",
                    original_name,
                    content_type,
                    len(raw_bytes),
                    gcs_raw_uri,
                    gcs_extracted_uri,
                    text_fingerprint,
                    extraction_ok,
                    extraction_method,
                    extracted_chars,
                ),
            )

        _log_job_event(conn, tenant_id, "doc_row_created", {"doc_id": doc_id, "fingerprint": fingerprint}, job_id=ingestion_job_id)
        _log_job_event(conn, tenant_id, "step_done", {"step": "ingest_file", "doc_id": doc_id}, job_id=ingestion_job_id)

    return IngestResponse(
        status="ok",
        kb_id=kb_id,
        doc_id=doc_id,
        ingestion_job_id=ingestion_job_id,
        fingerprint=fingerprint,
        text_fingerprint=text_fingerprint,
        scope=kb_scope_str,
        gcs_raw_uri=gcs_raw_uri,
        gcs_extracted_uri=gcs_extracted_uri,
        extraction_ok=extraction_ok,
        extraction_method=extraction_method,
        extracted_chars=extracted_chars,
        meta=extract_meta,
    )


# ---------------------------
# Real self-test (no mock data)
# ---------------------------
if __name__ == "__main__":
    """
    Real checks only:
    - DB connect + ensure documents table exists
    - GCS connect + bucket exists (list 1 object)
    No fake inserts/files unless you set TEST_INGEST_LOCAL_FILE.
    """
    settings = get_settings()

    # DB
    with _db_conn(settings) as conn:
        _ensure_documents_table(conn)
        with conn.cursor() as cur:
            cur.execute("SELECT 1;")
            print("DB OK:", cur.fetchone()[0])

    # GCS
    bucket_name = _pick(settings, "GCS_BUCKET_NAME")
    client = _gcs_client(settings)
    bucket = client.bucket(bucket_name)
    blobs = list(client.list_blobs(bucket, max_results=1))
    print("GCS OK:", bucket_name, "sample_blob_found:", bool(blobs))

    # Optional: ingest a real local file if provided
    local_path = os.getenv("TEST_INGEST_LOCAL_FILE")
    if local_path and os.path.exists(local_path):
        print("NOTE: TEST_INGEST_LOCAL_FILE is set. Use API endpoint for ingestion, not __main__ path.")
