"""
Shared markdown-to-Word converter used by both the download API endpoint
and the scripts/md_to_docx.py CLI tool.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Encoding cleanup ──────────────────────────────────────────────────────────
# Family A: UTF-8 3-byte sequences mis-decoded as Latin-1.
# Each Latin-1 code point == its original byte, so we reconstruct the UTF-8
# bytes and re-decode. Covers the full General Punctuation / Misc block.
_LAT1_RE = re.compile(
    'â'                  # original byte 0xE2  -> Latin-1 U+00E2
    '[-¿]'         # continuation byte 2 (0x80-0xBF)
    '[-¿]'         # continuation byte 3 (0x80-0xBF)
)

# Family B: same bytes mis-decoded as CP1252.
# Use single-quoted strings throughout; the content may contain curly quotes.
_CP1252 = [
    ('â€™', '’'),   # right single quote  '
    ('â€œ', '“'),   # left  double quote  "
    ('â€', '”'),   # right double quote  "
    ('â€“', '—'),   # em-dash             -
    ('â€–', '–'),   # en-dash             -
    ('â€¦', '…'),   # ellipsis            ...
    ('â€˜', '‘'),   # left  single quote  '
]


def _lat1_fix(m: re.Match) -> str:
    raw = bytes(ord(c) for c in m.group(0))
    try:
        return raw.decode('utf-8')
    except UnicodeDecodeError:
        return m.group(0)


def clean(text: str) -> str:
    """
    Fix all mojibake, then replace typographic dashes with plain hyphens
    so the output reads naturally without AI-style long dashes.
    """
    text = text.lstrip('﻿')                  # strip BOM
    text = _LAT1_RE.sub(_lat1_fix, text)          # family A: reconstruct
    for bad, good in _CP1252:                     # family B: literal swap
        text = text.replace(bad, good)
    # After mojibake fix, replace any remaining typographic dashes
    text = text.replace('—', ' - ')          # em-dash
    text = text.replace('–', ' - ')          # en-dash
    text = text.replace('‑', '-')            # non-breaking hyphen
    # Convert markdown links [text](url) → text (Word doesn't render MD links)
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    return text


# ── Inline markdown -> Word runs ──────────────────────────────────────────────
_INLINE = re.compile(
    r'\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|`(.+?)`'
)


def _add_inline(para, text: str) -> None:
    text = clean(text)
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            r = para.add_run(m.group(1)); r.bold = True; r.italic = True
        elif m.group(2):
            r = para.add_run(m.group(2)); r.bold = True
        elif m.group(3):
            r = para.add_run(m.group(3)); r.italic = True
        elif m.group(4):
            r = para.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r'^\|[\s\-:|]+\|', line.strip()))


def _parse_table_row(line: str) -> list:
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def _add_md_table(doc: Document, rows: list) -> None:
    """Render a list of row-lists as a proper Word table."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        tr = table.add_row()
        for j in range(n_cols):
            cell = tr.cells[j]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            cell_text = row_data[j] if j < len(row_data) else ''
            _add_inline(para, cell_text)
            if i == 0:
                for run in para.runs:
                    run.bold = True

    doc.add_paragraph('')


def _add_hr(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Public converter ───────────────────────────────────────────────────────────
def markdown_to_docx(
    title: str,
    markdown: str,
    dest: Union[Path, io.BytesIO],
) -> None:
    """
    Convert *markdown* to a Word document and write to *dest*.

    *dest* may be a filesystem Path or an in-memory BytesIO buffer.
    Seek the buffer to position 0 after this call before reading from it.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Paragraph spacing (applied once via styles — no per-paragraph overhead) ─
    doc.styles['Normal'].paragraph_format.space_after  = Pt(8)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)

    for _h, _before, _after in [
        ('Heading 1', 20, 8),
        ('Heading 2', 18, 6),
        ('Heading 3', 14, 4),
    ]:
        doc.styles[_h].paragraph_format.space_before = Pt(_before)
        doc.styles[_h].paragraph_format.space_after  = Pt(_after)

    doc.styles['List Bullet'].paragraph_format.space_after = Pt(3)
    doc.styles['List Number'].paragraph_format.space_after = Pt(3)

    tp = doc.add_heading(clean(title), level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(14)
    doc.add_paragraph('')

    table_buf: list = []

    def _flush_table():
        if table_buf:
            _add_md_table(doc, table_buf)
            table_buf.clear()

    for line in markdown.splitlines():
        stripped = line.strip()

        # Blank line — flush any buffered table, skip the line
        if not stripped:
            _flush_table()
            continue

        # Strip CMS metadata comments — they belong in the CMS, not the Word doc
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            continue

        # Markdown table row
        if stripped.startswith('|') and stripped.endswith('|'):
            if not _is_table_separator(stripped):
                table_buf.append(_parse_table_row(stripped))
            continue  # separator rows are skipped entirely

        # Any non-table line — flush buffered table first
        _flush_table()

        if stripped.startswith('### '):
            h = re.sub(r'\*+(.+?)\*+', r'\1', clean(stripped[4:]))
            h = re.sub(r'\s*\{#[^}]+\}', '', h)
            doc.add_heading(h, level=3)
        elif stripped.startswith('## '):
            h = re.sub(r'\*+(.+?)\*+', r'\1', clean(stripped[3:]))
            h = re.sub(r'\s*\{#[^}]+\}', '', h)
            doc.add_heading(h, level=2)
        elif stripped.startswith('# '):
            h = re.sub(r'\*+(.+?)\*+', r'\1', clean(stripped[2:]))
            h = re.sub(r'\s*\{#[^}]+\}', '', h)
            doc.add_heading(h, level=1)
        elif re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            _add_hr(doc)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_inline(doc.add_paragraph(style='List Bullet'), stripped[2:])
        elif re.match(r'^\d+[\.\)]\s', stripped):
            _add_inline(
                doc.add_paragraph(style='List Number'),
                re.sub(r'^\d+[\.\)]\s+', '', stripped),
            )
        elif stripped.startswith('> '):
            _add_inline(doc.add_paragraph(style='Quote'), stripped[2:])
        else:
            _add_inline(doc.add_paragraph(), stripped)

    # Flush any table that ends at the last line
    _flush_table()

    if isinstance(dest, Path):
        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest))
    else:
        doc.save(dest)
